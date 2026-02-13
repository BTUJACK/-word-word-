import os
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class DocxTableAdder:
    def __init__(self, root):
        self.root = root
        self.root.title("批量添加Word表格工具（ME/RE区分版）")
        self.root.geometry("700x550")
        
        # 选择文件夹相关
        self.folder_path = tk.StringVar()
        
        # 创建GUI组件
        self._create_widgets()
    
    def _create_widgets(self):
        # 文件夹选择区域
        frame1 = tk.Frame(self.root, padx=10, pady=10)
        frame1.pack(fill=tk.X)
        
        tk.Label(frame1, text="目标文件夹:").pack(side=tk.LEFT)
        tk.Entry(frame1, textvariable=self.folder_path, width=50).pack(side=tk.LEFT, padx=5)
        tk.Button(frame1, text="浏览", command=self._select_folder).pack(side=tk.LEFT)
        
        # 功能说明区域
        frame_info = tk.Frame(self.root, padx=10, pady=5, bg="#f0f0f0")
        frame_info.pack(fill=tk.X)
        
        tk.Label(frame_info, text="功能说明：", bg="#f0f0f0", font=("Arial", 9, "bold")).pack(anchor=tk.W)
        info_text = "• 包含\"ME\"的文件：插入频率范围150kHz-30MHz的表格\n• 包含\"RE\"的文件：插入频率范围30MHz-1GHz的表格\n• 不包含关键词的文件：跳过处理"
        tk.Label(frame_info, text=info_text, bg="#f0f0f0", font=("Arial", 9), justify=tk.LEFT).pack(anchor=tk.W)
        
        # 操作按钮区域
        frame2 = tk.Frame(self.root, padx=10, pady=5)
        frame2.pack(fill=tk.X)
        
        tk.Button(frame2, text="开始处理", command=self._process_files, bg="#4CAF50", fg="white").pack(side=tk.LEFT, padx=5)
        tk.Button(frame2, text="清空日志", command=self._clear_log, bg="#f44336", fg="white").pack(side=tk.LEFT, padx=5)
        
        # 日志显示区域
        frame3 = tk.Frame(self.root, padx=10, pady=10)
        frame3.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(frame3, text="处理日志:").pack(anchor=tk.W)
        self.log_text = scrolledtext.ScrolledText(frame3, height=20)
        self.log_text.pack(fill=tk.BOTH, expand=True)
    
    def _select_folder(self):
        """选择目标文件夹"""
        folder = filedialog.askdirectory(title="选择包含docx文件的文件夹")
        if folder:
            self.folder_path.set(folder)
            self._log(f"已选择文件夹: {folder}")
    
    def _log(self, message):
        """添加日志信息"""
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
    
    def _clear_log(self):
        """清空日志"""
        self.log_text.delete(1.0, tk.END)
    
    def _extract_doc_text(self, doc):
        """
        提取docx文档的所有文本内容（用于关键词检测）
        :param doc: Document对象
        :return: 文档纯文本字符串
        """
        full_text = []
        # 提取段落文本
        for para in doc.paragraphs:
            full_text.append(para.text)
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return " ".join(full_text)
    
    def _detect_keyword(self, file_path):
        """
        检测文档中是否包含ME或RE关键词
        :param file_path: 文件路径
        :return: "ME" / "RE" / None
        """
        try:
            doc = Document(file_path)
            text = self._extract_doc_text(doc).upper()  # 转为大写，不区分大小写
            if "ME" in text:
                return "ME"
            elif "RE" in text:
                return "RE"
            else:
                return None
        except Exception as e:
            self._log(f"检测文件 {file_path} 关键词出错: {str(e)}")
            return None
    
    def _set_cell_border(self, cell, border_color="000000", border_width=1):
        """
        为表格单元格设置边框（不依赖样式，兼容性100%）
        :param cell: 单元格对象
        :param border_color: 边框颜色（16进制）
        :param border_width: 边框宽度（点数）
        """
        # 创建边框元素
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = parse_xml(r'<w:tcBorders %s>'
                               r'<w:top w:val="single" w:color="%s" w:sz="%d"/>'
                               r'<w:left w:val="single" w:color="%s" w:sz="%d"/>'
                               r'<w:bottom w:val="single" w:color="%s" w:sz="%d"/>'
                               r'<w:right w:val="single" w:color="%s" w:sz="%d"/>'
                               r'</w:tcBorders>' % (nsdecls('w'),
                                                     border_color, border_width*8,
                                                     border_color, border_width*8,
                                                     border_color, border_width*8,
                                                     border_color, border_width*8))
        # 添加边框到单元格
        tc_pr.append(tc_borders)
    
    def _add_table_to_docx(self, file_path, keyword_type):
        """
        根据关键词类型给单个docx文件添加对应表格（完整保留所有内容，包括图片）
        :param file_path: 文件路径
        :param keyword_type: "ME" / "RE"
        :return: 处理结果 True/False
        """
        try:
            # 打开文档（保留所有原始内容）
            doc = Document(file_path)
            
            # ========== 根据关键词选择表格内容 ==========
            if keyword_type == "ME":
                frequency_range = "150kHz-30MHz"
            elif keyword_type == "RE":
                frequency_range = "30MHz-1GHz"
            else:
                return False
            
            # ========== 在文档开头插入表格 ==========
            # 步骤1: 在文档最开头插入一个空段落作为占位
            body = doc.element.body
            new_para = parse_xml(r'<w:p %s></w:p>' % nsdecls('w'))
            body.insert(0, new_para)
            
            # 步骤2: 创建2行2列的表格
            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # 设置表格内容
            table.cell(0, 0).text = "试验供电电源：380V AC/50Hz"
            table.cell(0, 1).text = f"试验频率范围：{frequency_range}"
            table.cell(1, 0).text = "样品运行模式：1"
            table.cell(1, 1).text = ""  # 第二行第二列留空
            
            # 为表格添加边框
            for row in table.rows:
                for cell in row.cells:
                    self._set_cell_border(cell)
                    # 设置单元格宽度和字体
                    cell.width = Inches(2.5)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            
            # 步骤3: 将表格移动到文档最开头
            table_elem = table._element
            table_elem.getparent().remove(table_elem)
            body.insert(1, table_elem)
            
            # 步骤4: 在表格后添加两行空行
            empty_para1 = parse_xml(r'<w:p %s><w:r><w:t></w:t></w:r></w:p>' % nsdecls('w'))
            empty_para2 = parse_xml(r'<w:p %s><w:r><w:t></w:t></w:r></w:p>' % nsdecls('w'))
            body.insert(2, empty_para1)
            body.insert(3, empty_para2)
            
            # 删除临时空段落
            body.remove(new_para)
            
            # 保存修改后的文档
            doc.save(file_path)
            return True
            
        except Exception as e:
            self._log(f"处理文件 {file_path} 出错: {str(e)}")
            import traceback
            self._log(f"详细错误信息: {traceback.format_exc()}")
            return False
    
    def _process_files(self):
        """批量处理文件夹中的docx文件（按ME/RE关键词区分）"""
        folder = self.folder_path.get()
        if not folder or not os.path.exists(folder):
            messagebox.showerror("错误", "请选择有效的文件夹！")
            return
        
        # 获取所有docx文件
        docx_files = [f for f in os.listdir(folder) if f.lower().endswith('.docx')]
        if not docx_files:
            messagebox.showinfo("提示", "文件夹中未找到docx文件！")
            return
        
        self._log(f"找到 {len(docx_files)} 个docx文件，开始处理...")
        
        # 统计变量
        me_count = 0       # ME文件处理数
        re_count = 0       # RE文件处理数
        skip_count = 0     # 跳过文件数
        fail_count = 0     # 失败文件数
        
        for filename in docx_files:
            file_path = os.path.join(folder, filename)
            self._log(f"\n正在处理: {filename}")
            
            # 第一步：检测关键词
            keyword = self._detect_keyword(file_path)
            if keyword is None:
                self._log(f"文件 {filename} 未检测到ME/RE关键词，跳过处理")
                skip_count += 1
                continue
            
            # 第二步：创建备份文件
            backup_path = file_path + ".bak"
            try:
                shutil.copy2(file_path, backup_path)
                self._log(f"已创建备份: {backup_path}")
            except Exception as e:
                self._log(f"创建备份失败 {filename}: {str(e)}")
                fail_count += 1
                continue
            
            # 第三步：根据关键词处理文件
            if self._add_table_to_docx(file_path, keyword):
                self._log(f"成功处理【{keyword}类型】文件: {filename}")
                if keyword == "ME":
                    me_count += 1
                else:
                    re_count += 1
            else:
                # 处理失败则恢复备份
                try:
                    shutil.copy2(backup_path, file_path)
                    self._log(f"恢复备份: {filename}")
                except:
                    pass
                fail_count += 1
        
        # 处理完成统计
        total_process = me_count + re_count
        self._log(f"\n========== 处理完成 ==========")
        self._log(f"ME类型文件处理成功: {me_count} 个")
        self._log(f"RE类型文件处理成功: {re_count} 个")
        self._log(f"跳过无关键词文件: {skip_count} 个")
        self._log(f"处理失败文件: {fail_count} 个")
        self._log(f"总计处理文件: {total_process} 个")
        
        # 弹窗提示
        messagebox.showinfo("处理完成", 
                           f"处理结果汇总：\n"
                           f"ME类型文件：{me_count} 个（已插入150kHz-30MHz表格）\n"
                           f"RE类型文件：{re_count} 个（已插入30MHz-1GHz表格）\n"
                           f"跳过无关键词文件：{skip_count} 个\n"
                           f"处理失败文件：{fail_count} 个")

if __name__ == "__main__":
    # 安装依赖提示（首次运行需要）
    try:
        from docx import Document
    except ImportError:
        root = tk.Tk()
        root.withdraw()
        messagebox.showinfo("提示", "请先安装依赖库：\npip install python-docx")
        os.system("pip install python-docx")
    
    # 启动GUI
    root = tk.Tk()
    app = DocxTableAdder(root)
    root.mainloop()
