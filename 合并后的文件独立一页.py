#合并后的文档中每个源 Word 文档的内容独立占一页（而非所有内容连续紧挨着）
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import os
import datetime
from docx import Document
from docxcompose.composer import Composer
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_BREAK  # 关键：导入分页符枚举类

# 安装依赖（执行以下命令）：
# pip install python-docx python-docx-composer

class WordMergerGUI:
    def __init__(self, root):
        # 主窗口配置
        self.root = root
        self.root.title("Word文档合并工具（每页独立保留源文档内容）")
        self.root.geometry("700x400")
        self.root.resizable(False, False)
        
        # ========== 1. 文件夹选择区域 ==========
        frame1 = tk.Frame(root, padx=20, pady=15)
        frame1.pack(fill=tk.X)
        
        tk.Label(frame1, text="待合并Word文件夹：", font=("微软雅黑", 11)).grid(row=0, column=0, sticky=tk.W)
        self.folder_var = tk.StringVar()
        entry_folder = tk.Entry(frame1, textvariable=self.folder_var, width=40, font=("微软雅黑", 10))
        entry_folder.grid(row=0, column=1, padx=10)
        
        btn_folder = tk.Button(
            frame1, text="选择文件夹", command=self.select_folder,
            font=("微软雅黑", 10), bg="#1E90FF", fg="white", width=10
        )
        btn_folder.grid(row=0, column=2)
        
        # ========== 2. 输出路径选择区域 ==========
        frame2 = tk.Frame(root, padx=20, pady=5)
        frame2.pack(fill=tk.X)
        
        tk.Label(frame2, text="合并后保存路径：", font=("微软雅黑", 11)).grid(row=0, column=0, sticky=tk.W)
        self.output_var = tk.StringVar()
        default_out = os.path.join(os.getcwd(), f"merged_{datetime.datetime.now().strftime('%Y%m%d')}.docx")
        self.output_var.set(default_out)
        
        entry_output = tk.Entry(frame2, textvariable=self.output_var, width=40, font=("微软雅黑", 10))
        entry_output.grid(row=0, column=1, padx=10)
        
        btn_output = tk.Button(
            frame2, text="选择保存位置", command=self.select_output,
            font=("微软雅黑", 10), bg="#1E90FF", fg="white", width=10
        )
        btn_output.grid(row=0, column=2)
        
        # ========== 3. 合并按钮 ==========
        frame3 = tk.Frame(root, padx=20, pady=20)
        frame3.pack(fill=tk.X)
        
        self.btn_merge = tk.Button(
            frame3, text="开始合并文档（每页独立）", command=self.merge_documents,
            font=("微软雅黑", 14, "bold"), bg="#32CD32", fg="white",
            width=20, height=2
        )
        self.btn_merge.pack()
        
        # ========== 4. 日志显示区域 ==========
        frame4 = tk.Frame(root, padx=20, pady=5)
        frame4.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(frame4, text="操作日志：", font=("微软雅黑", 10)).pack(anchor=tk.W)
        self.log_text = scrolledtext.ScrolledText(frame4, height=8, font=("Consolas", 9))
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        # 初始化日志
        self.log("✅ 工具已就绪，选择文件夹后点击合并即可（每页独立保留源文档内容）")

    # ========== 辅助方法 ==========
    def log(self, content):
        """添加带时间戳的日志"""
        time_str = datetime.datetime.now().strftime("[%Y-%m-%d %H:%M:%S]")
        self.log_text.insert(tk.END, f"{time_str} {content}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def select_folder(self):
        """选择待合并的Word文件夹"""
        folder = filedialog.askdirectory(title="选择存放多个Word文件的文件夹")
        if folder:
            self.folder_var.set(folder)
            docx_count = len([f for f in os.listdir(folder) if f.lower().endswith(".docx")])
            self.log(f"📂 已选择文件夹：{folder}")
            self.log(f"🔍 检测到 {docx_count} 个.docx文件待合并")

    def select_output(self):
        """选择输出路径"""
        file = filedialog.asksaveasfilename(
            title="选择合并后文件的保存位置",
            defaultextension=".docx",
            filetypes=[("Word 2007-2019文档", "*.docx"), ("所有文件", "*.*")]
        )
        if file:
            self.output_var.set(file)
            self.log(f"💾 已选择输出路径：{file}")

    # ========== 核心合并方法（修复分页符参数错误，每页独立保留源文档内容） ==========
    def merge_documents(self):
        """合并Word，每个源文档独立占一页（添加分页+分节符）"""
        try:
            # 1. 获取输入路径
            source_folder = self.folder_var.get().strip()
            output_path = self.output_var.get().strip()
            
            # 2. 校验输入
            if not source_folder or not os.path.exists(source_folder):
                messagebox.showerror("错误", "请选择有效的待合并文件夹！")
                return
            if not output_path:
                messagebox.showerror("错误", "请选择合并后文件的保存路径！")
                return
            
            # 3. 筛选docx文件
            docx_files = [
                os.path.join(source_folder, f)
                for f in os.listdir(source_folder)
                if f.lower().endswith(".docx") and os.path.isfile(os.path.join(source_folder, f))
            ]
            
            if not docx_files:
                messagebox.showwarning("警告", "所选文件夹内无有效的.docx文件！")
                return
            
            self.log("="*50)
            self.log(f"🚀 开始合并 - 共 {len(docx_files)} 个文件（每页独立）")
            self.log("="*50)
            
            # 4. 核心合并逻辑（添加分节符+分页符，确保每页独立）
            # 以第一个文档为基础
            master_doc = Document(docx_files[0])
            composer = Composer(master_doc)
            
            # 逐个追加其他文档（每个文档前加分节符+分页符）
            for idx, file_path in enumerate(docx_files[1:], 2):
                self.log(f"📄 正在合并第 {idx} 个文件：{os.path.basename(file_path)}")
                
                # 打开当前文档
                doc = Document(file_path)
                
                # 关键：修复后的分页符插入方式（使用WD_BREAK.PAGE）
                doc.paragraphs[0].insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
                
                # 设置节的起始位置为新页（双重保障）
                for section in doc.sections:
                    section.start_type = WD_SECTION_START.NEW_PAGE
                
                # 追加文档（此时会自动从新页开始）
                composer.append(doc)
            
            # 5. 保存合并后的文档
            composer.save(output_path)
            
            # 6. 合并完成
            self.log("="*50)
            self.log(f"🎉 合并成功！每个源文档独立占一页")
            self.log(f"📁 输出文件：{output_path}")
            self.log("="*50)
            
            messagebox.showinfo("合并完成", 
                f"✅ 文档合并成功！\n"
                f"📄 共合并 {len(docx_files)} 个Word文件\n"
                f"📄 每个源文档内容独立保留在一页\n"
                f"💾 输出路径：\n{output_path}")
        
        except Exception as e:
            self.log(f"❌ 合并失败：{str(e)}")
            messagebox.showerror("合并失败", f"合并过程出错：\n{str(e)}")

if __name__ == "__main__":
    # 适配Windows高分屏
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except:
        pass
    
    # 启动主窗口
    root = tk.Tk()
    app = WordMergerGUI(root)
    root.mainloop()
