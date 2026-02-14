#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量修改docx表格工具 - Python 3.8.7 + python-docx 0.8.11
功能：
1. 按文件名关键词（ME_H/ME_V/RE_H/RE_V）差异化处理表格
2. 第三列右侧加3列并填充对应内容，原第四列移第七列
3. 添加合并列的第八行备注，显示完整表格边框




用Python 3.8.7实现批量修改一个文件夹里面的.docx文件，并且生成一个GUI界面进行操作。
1
如果.docx文件含有“ME_H”，在.docx文件表格第三列右侧添加三列表格并且填充内容：
天线高度(cm)    天线极化    转台角度(deg)
130         H       ——
130         H       ——
130         H       ——
130         H       ——
130         H       ——
130         H       ——

原来的第四列表格和内容放在第7列
添加第八行，所有列合并为一列，添加内容并且靠左对齐：
备注：——
2
如果.docx文件含有“ME_V”，在.docx文件表格第三列右侧添加三列表格并且填充内容：
天线高度(cm)    天线极化    转台角度(deg)
130         V       ——
130         V       ——
130         V       ——
130         V       ——
130         V       ——
130         V       ——

原来的第四列表格和内容放在第7列
添加第八行，所有列合并为一列，添加内容并且靠左对齐：
备注：——
3
如果.docx文件含有“RE_H”，在.docx文件表格第三列右侧添加三列表格并且填充内容：
天线高度(cm)    天线极化    转台角度(deg)
200         H       ——
200         H       ——
200         H       ——
200         H       ——
200         H       ——
200         H       ——

原来的第四列表格和内容放在第7列
添加第八行，所有列合并为一列，添加内容并且靠左对齐：
备注：背景噪声超限值频段除外，其余频段峰值均低于限值

4
如果.docx文件含有“RE_V”，在.docx文件表格第三列右侧添加三列表格并且填充内容：
天线高度(cm)    天线极化    转台角度(deg)
200         H       ——
200         H       ——
200         H       ——
200         H       ——
200         H       ——
200         H       ——

原来的第四列表格和内容放在第7列
添加第八行，所有列合并为一列，添加内容并且靠左对齐：
备注：背景噪声超限值频段除外，其余频段峰值均低于限值

豆包话术：
合并文档之前，对每个文档进行操作：
删除图片以上的所有内容，包括页眉；
将表格整体移动到图片上面；
删除“Final_Result”；
图片左上角的上面一行添加文字“试验结果图:”；
如果文档标题包含字母：“H”，就在图片正下方居中添加文字：“水平极化”；
如果文档标题包含字母：“V”，就在图片正下方居中添加文字：“垂直极化”；
删除“Margin”列右侧的所有表格和内容；
“Margin”列全部内容移动到“Limit”列左侧；
“Frequency”替换为“频率”；
“QuasiPeak”替换为“准峰值”；
“Margin”替换为“裕量”；
“Limit”替换为“限值”；
合并文档之后，恢复每个文档以前的内容。


"""
import os
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocxTableModifier:
    def __init__(self, root):
        self.root = root
        self.root.title("Docx表格批量修改工具（关键词差异化处理）")
        self.root.geometry("750x550")
        
        # 配置不同关键词对应的参数
        self.config = {
            "ME_H": {
                "data_values": ["130", "H", "——"],
                "remark": "备注：——"
            },
            "ME_V": {
                "data_values": ["130", "V", "——"],
                "remark": "备注：——"
            },
            "RE_H": {
                "data_values": ["200", "H", "——"],
                "remark": "备注：背景噪声超限值频段除外，其余频段峰值均低于限值"
            },
            "RE_V": {
                "data_values": ["200", "H", "——"],
                "remark": "备注：背景噪声超限值频段除外，其余频段峰值均低于限值"
            }
        }
        self.default_config = {
            "data_values": ["", "", ""],
            "remark": "备注：无匹配关键词"
        }
        
        self.folder_path = tk.StringVar()
        self._create_gui()

    def _create_gui(self):
        """创建GUI界面"""
        # 1. 文件夹选择区域
        frame1 = tk.Frame(self.root, padx=10, pady=10)
        frame1.pack(fill=tk.X)
        
        tk.Label(frame1, text="目标文件夹：", font=("SimHei", 10)).pack(side=tk.LEFT)
        tk.Entry(frame1, textvariable=self.folder_path, width=55, font=("SimHei", 10)).pack(side=tk.LEFT, padx=5)
        tk.Button(
            frame1, text="选择文件夹", command=self._select_folder,
            font=("SimHei", 10), bg="#E0E0E0"
        ).pack(side=tk.LEFT)
        
        # 2. 执行按钮
        frame2 = tk.Frame(self.root, padx=10, pady=5)
        frame2.pack(fill=tk.X)
        
        tk.Button(
            frame2, text="开始批量处理", 
            command=self._batch_process,
            bg="#2196F3", fg="white", font=("SimHei", 11, "bold"), padx=20
        ).pack(side=tk.LEFT)
        
        # 3. 日志显示区域
        frame3 = tk.Frame(self.root, padx=10, pady=10)
        frame3.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(frame3, text="处理日志：", font=("SimHei", 10)).pack(anchor=tk.W)
        self.log_text = scrolledtext.ScrolledText(
            frame3, height=22, font=("Consolas", 9), wrap=tk.WORD
        )
        self.log_text.pack(fill=tk.BOTH, expand=True)

    def _select_folder(self):
        """选择目标文件夹"""
        folder = filedialog.askdirectory(title="选择包含docx文件的文件夹")
        if folder:
            self.folder_path.set(folder)
            self._log(f"✅ 已选择文件夹：{folder}")

    def _log(self, msg):
        """添加日志信息并自动滚动"""
        self.log_text.insert(tk.END, f"{msg}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def _set_cell_border(self, cell):
        """为单元格设置完整边框（黑色0.5磅实线）"""
        borders = ["top", "bottom", "left", "right"]
        for border_name in borders:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")       # 实线边框
            border.set(qn("w:sz"), "4")             # 0.5磅宽度（1/8磅单位）
            border.set(qn("w:color"), "000000")     # 黑色
            border.set(qn("w:space"), "0")          # 无间距
            cell._tc.get_or_add_tcPr().append(border)

    def _get_file_config(self, file_name):
        """根据文件名匹配配置"""
        for keyword in self.config.keys():
            if keyword in file_name:
                return self.config[keyword]
        return self.default_config

    def _rebuild_table(self, table, data_values):
        """重建表格数据：原1-3列+新增3列+原4列"""
        # 1. 读取原表格内容
        original_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            original_data.append(row_data)
        
        if not original_data:
            return None
        
        # 2. 构造新表格数据
        new_table_data = []
        new_col_headers = ["天线高度(cm)", "天线极化", "转台角度(deg)"]
        for idx, row in enumerate(original_data):
            # 补全原行数据至4列
            row += [""] * (4 - len(row))
            
            # 表头行填新增列标题，数据行填对应值（前6行填指定值）
            if idx == 0:
                new_cols = new_col_headers
            elif 1 <= idx <= 6:  # 第2-7行（数据行）填配置值
                new_cols = data_values
            else:
                new_cols = ["", "", ""]  # 超出6行填空
            
            # 新行结构：原1-3列 + 新增3列 + 原4列（第7列）
            new_row = row[0:3] + new_cols + [row[3]]
            new_table_data.append(new_row)
        
        return new_table_data

    def _add_remark_row(self, table, remark_text):
        """为表格添加第八行（合并所有列），填入备注"""
        # 添加新行（第八行）
        new_row = table.add_row().cells
        col_count = len(table.columns)
        
        # 合并所有列
        for i in range(1, col_count):
            new_row[0].merge(new_row[i])
        
        # 设置单元格内容和格式
        cell = new_row[0]
        cell.text = remark_text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直居中
        # 设置文字居中
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 为合并后的单元格设置边框
        self._set_cell_border(cell)

    def _modify_docx_table(self, file_path, file_config):
        """修改单个docx文件的表格"""
        # 1. 备份原文件
        backup_path = f"{file_path}.bak"
        shutil.copy2(file_path, backup_path)
        self._log(f"  📁 已备份原文件：{os.path.basename(backup_path)}")
        
        # 2. 打开文档处理表格
        doc = Document(file_path)
        table_count = 0
        
        for table in doc.tables:
            table_count += 1
            self._log(f"  📋 处理第{table_count}个表格（原行列数：{len(table.rows)}行 × {len(table.columns)}列）")
            
            # 跳过列数不足4的表格
            if len(table.columns) < 4:
                self._log(f"  ⚠️  第{table_count}个表格列数不足4列，跳过")
                continue
            
            # 3. 重建表格数据
            new_table_data = self._rebuild_table(table, file_config["data_values"])
            if not new_table_data:
                self._log(f"  ⚠️  第{table_count}个表格无数据，跳过")
                continue
            
            # 4. 删除原表格
            table_element = table._element
            table_parent = table_element.getparent()
            table_idx = list(table_parent).index(table_element)
            table_parent.remove(table_element)
            
            # 5. 创建新表格并填充数据
            new_table = doc.add_table(rows=len(new_table_data), cols=7)
            new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 设置列宽
            for col in new_table.columns:
                col.width = Pt(60)
            
            # 填充数据+设置边框
            for row_idx, row_data in enumerate(new_table_data):
                row_cells = new_table.rows[row_idx].cells
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(row_cells):
                        cell = row_cells[col_idx]
                        cell.text = cell_text
                        self._set_cell_border(cell)
            
            # 6. 添加第八行备注（合并列）
            self._add_remark_row(new_table, file_config["remark"])
            
            # 7. 将新表格插入原位置
            table_parent.insert(table_idx, new_table._element)
        
        # 8. 保存文档
        doc.save(file_path)
        self._log(f"  ✅ 已完成文件修改：{os.path.basename(file_path)}")
        return True

    def _batch_process(self):
        """批量处理文件夹下的docx文件"""
        folder = self.folder_path.get()
        if not folder or not os.path.exists(folder):
            messagebox.showerror("错误", "请选择有效的文件夹！")
            return
        
        # 清空日志
        self.log_text.delete(1.0, tk.END)
        self._log("🚀 开始批量处理docx文件...")
        
        # 获取所有docx文件
        docx_files = [
            f for f in os.listdir(folder)
            if f.lower().endswith(".docx") and os.path.isfile(os.path.join(folder, f))
        ]
        
        if not docx_files:
            self._log("⚠️  未找到任何.docx文件！")
            messagebox.showinfo("提示", "未找到任何.docx文件！")
            return
        
        self._log(f"📊 共找到 {len(docx_files)} 个docx文件，开始处理...")
        
        # 批量处理
        success = 0
        fail = 0
        for file_name in docx_files:
            file_path = os.path.join(folder, file_name)
            self._log(f"\n🔍 处理文件：{file_name}")
            
            # 获取当前文件的配置
            file_config = self._get_file_config(file_name)
            self._log(f"  📌 匹配关键词：{[k for k in self.config if k in file_name] or '无'}")
            
            try:
                self._modify_docx_table(file_path, file_config)
                success += 1
            except Exception as e:
                self._log(f"❌ 处理失败：{str(e)}")
                fail += 1
        
        # 处理完成提示
        result = f"✅ 处理完成！成功：{success}个 | 失败：{fail}个"
        self._log(f"\n{result}")
        messagebox.showinfo("完成", result)

if __name__ == "__main__":
    # 适配中文显示
    root = tk.Tk()
    root.option_add("*Font", "SimHei 9")
    app = DocxTableModifier(root)
    root.mainloop()
