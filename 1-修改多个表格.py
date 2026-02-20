#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合并版Docx批量处理工具：先重命名文件，再处理文档内容
1. 重命名：将M1_/M2_/M3_/M4_/M5_/Ambient_移动到P1_后（直接修改原文件，不保留副本）
2. 内容处理：删除指定文本、调整表格、扩展表格、图片标注等（直接覆盖原文件，不备份）
"""


'''

用Python 3.8.7实现把一个文件夹里面的.docx文件批量修改名字, 并且生成一个GUI界面进行操作：
“M1_”移动到“P1_”后面；
“M2_”移动到“P1_”后面；
“M3_”移动到“P1_”后面；
“M4_”移动到“P1_”后面；
“M5_”移动到“P1_”后面；
“Ambient_”移动到“P1_”后面；
'''

'''
用Python 3.8.7实现批量修改一个文件夹里面的.docx文件，并且生成一个GUI界面进行操作。
删除“Test Report”；
删除每个文档中的第一个表格；
删除“Final_Result”
“Frequency”替换为“频率”；
“QuasiPeak”替换为“准峰值”；
“Margin”替换为“裕量”；
“Limit”替换为“限值”；
删除表格的第5列到第9列；
交换表格第3列和第4列的内容；




用Python 3.8.7实现批量修改一个文件夹里面的.docx文件，并且生成一个GUI界面进行操作。
1
如果.docx文件含有“ME_H”，在.docx文件表格第三列右侧添加三列表格并且填充内容：
天线高度(cm)    天线极化    转台角度(deg)
130         H       ——
130         H       ——
130         H       ——
130         H       ——
130         H       ——
130         H       ——

原来的第四列表格和内容放在第7列
添加第八行，所有列合并为一列，添加内容并且靠左对齐：
备注：——
2
如果.docx文件含有“ME_V”，在.docx文件表格第三列右侧添加三列表格并且填充内容：
天线高度(cm)    天线极化    转台角度(deg)
130         V       ——
130         V       ——
130         V       ——
130         V       ——
130         V       ——
130         V       ——

原来的第四列表格和内容放在第7列
添加第八行，所有列合并为一列，添加内容并且靠左对齐：
备注：——
3
如果.docx文件含有“RE_H”，在.docx文件表格第三列右侧添加三列表格并且填充内容：
天线高度(cm)    天线极化    转台角度(deg)
200         H       ——
200         H       ——
200         H       ——
200         H       ——
200         H       ——
200         H       ——

原来的第四列表格和内容放在第7列
添加第八行，所有列合并为一列，添加内容并且靠左对齐：
备注：背景噪声超限值频段除外，其余频段峰值均低于限值

4
如果.docx文件含有“RE_V”，在.docx文件表格第三列右侧添加三列表格并且填充内容：
天线高度(cm)    天线极化    转台角度(deg)
200         V       ——
200         V       ——
200         V       ——
200         V       ——
200         V       ——
200         V       ——

原来的第四列表格和内容放在第7列
添加第八行，所有列合并为一列，添加内容并且靠左对齐：
备注：背景噪声超限值频段除外，其余频段峰值均低于限值

'''


import os
import sys
import shutil
import traceback
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime

# 版本校验：确保使用Python 3.8及以上
assert sys.version_info >= (3, 8), "请使用Python 3.8及以上版本运行此程序"

class DocxBatchProcessor:
    def __init__(self, root):
        self.root = root
        self.root.title("Docx批量处理工具（直接修改原文件）")
        self.root.geometry("950x700")  # 加宽窗口，避免控件挤压
        self.root.resizable(True, True)  # 允许窗口缩放
        
        # 重命名配置
        self.target_prefixes = ["M1_", "M2_", "M3_", "M4_", "M5_", "Ambient_"]
        self.target_key = "P1_"
        
        # 内容处理配置
        self.table_config = {
            "ME_H": {
                "table_data": ["130", "H", "——"],
                "remark": "备注：——",
                "polarization": "水平极化"
            },
            "ME_V": {
                "table_data": ["130", "V", "——"],
                "remark": "备注：——",
                "polarization": "垂直极化"
            },
            "RE_H": {
                "table_data": ["200", "H", "——"],
                "remark": "备注：背景噪声超限值频段除外，其余频段峰值均低于限值",
                "polarization": "水平极化"
            },
            "RE_V": {
                "table_data": ["200", "V", "——"],
                "remark": "备注：背景噪声超限值频段除外，其余频段峰值均低于限值",
                "polarization": "垂直极化"
            }
        }
        self.default_config = {
            "table_data": ["", "", ""],
            "remark": "备注：无匹配关键词",
            "polarization": "未匹配极化类型"
        }
        
        # 界面变量
        self.source_folder = tk.StringVar()    # 源文件目录（直接处理此目录下的文件）
        self.log_text = None                   # 日志文本框
        
        # 创建GUI界面（改用grid布局，更可控）
        self._create_gui()

    def _create_gui(self):
        """创建修改后的GUI界面（移除输出文件夹选择，直接处理源文件夹）"""
        # 整体容器，添加统一内边距
        main_container = tk.Frame(self.root, padx=10, pady=10)
        main_container.pack(fill=tk.BOTH, expand=True)
        
        # 1. 文件夹选择区域（仅保留源文件夹）
        folder_frame = tk.LabelFrame(main_container, text="文件夹配置", padx=8, pady=8)
        folder_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 源文件夹行
        tk.Label(folder_frame, text="源文件夹：", font=("SimHei", 10), width=10).grid(
            row=0, column=0, sticky=tk.W, padx=(0, 5)
        )
        source_entry = tk.Entry(folder_frame, textvariable=self.source_folder, font=("SimHei", 10))
        source_entry.grid(row=0, column=1, sticky=tk.EW, padx=(0, 5))
        tk.Button(
            folder_frame, text="选择源文件夹", command=self._select_source_folder,
            font=("SimHei", 10), bg="#E0E0E0", width=12
        ).grid(row=0, column=2, padx=(0, 15))
        
        # 设置列权重，让输入框自适应宽度
        folder_frame.columnconfigure(1, weight=1)
        
        # 2. 执行按钮区域
        btn_frame = tk.Frame(main_container)
        btn_frame.pack(fill=tk.X, pady=(0, 10))
        
        tk.Button(
            btn_frame, text="开始完整处理（直接修改原文件）", 
            command=self._batch_process_all,
            bg="#2196F3", fg="white", font=("SimHei", 11, "bold"), padx=20, height=1
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        tk.Button(
            btn_frame, text="清空日志", 
            command=self._clear_log,
            bg="#f44336", fg="white", font=("SimHei", 10), padx=10, height=1
        ).pack(side=tk.LEFT)
        
        # 3. 警告提示
        warn_frame = tk.Frame(main_container)
        warn_frame.pack(fill=tk.X, pady=(0, 10))
        warn_label = tk.Label(
            warn_frame, 
            text="⚠️ 警告：此操作会直接修改源文件夹中的文件，不会保留原文件备份！请提前做好数据备份",
            font=("SimHei", 9), fg="#ff0000", wraplength=900
        )
        warn_label.pack(anchor=tk.W)
        
        # 4. 功能说明
        desc_frame = tk.Frame(main_container)
        desc_frame.pack(fill=tk.X, pady=(0, 10))
        desc_label = tk.Label(
            desc_frame, 
            text="功能1：重命名 - 将M1_/M2_/M3_/M4_/M5_/Ambient_移动到P1_后 | 功能2：内容修改 - 删除文本/调整表格/图片标注",
            font=("SimHei", 9), fg="#666666", wraplength=900  # 自动换行
        )
        desc_label.pack(anchor=tk.W)
        
        # 5. 日志显示区域（占满剩余空间）
        log_frame = tk.LabelFrame(main_container, text="处理日志", padx=8, pady=8)
        log_frame.pack(fill=tk.BOTH, expand=True)
        
        self.log_text = scrolledtext.ScrolledText(
            log_frame, font=("Consolas", 9), wrap=tk.WORD
        )
        self.log_text.pack(fill=tk.BOTH, expand=True)

    def _select_source_folder(self):
        """选择源文件夹"""
        folder = filedialog.askdirectory(title="选择要处理的docx文件所在文件夹（直接修改此目录文件）")
        if folder:
            self.source_folder.set(folder)
            self._log(f"✅ 已选择源文件夹：{folder}")

    def _log(self, msg):
        """日志输出（自动滚动）"""
        if self.log_text:
            self.log_text.insert(tk.END, f"{msg}\n")
            self.log_text.see(tk.END)
            self.root.update_idletasks()

    def _clear_log(self):
        """清空日志"""
        if self.log_text:
            self.log_text.delete(1.0, tk.END)
            self._log("📝 日志已清空")

    # -------------------------- 第一部分：重命名功能（直接修改原文件） --------------------------
    def _get_new_filename(self, old_name):
        """生成新文件名"""
        name_without_ext = os.path.splitext(old_name)[0]
        ext = os.path.splitext(old_name)[1]
        
        new_name = name_without_ext
        has_changed = False
        
        # 遍历需要移动的前缀
        for prefix in self.target_prefixes:
            if prefix in new_name and self.target_key in new_name:
                # 移除目标前缀
                new_name = new_name.replace(prefix, "")
                # 将目标前缀插入到P1_后面
                p1_index = new_name.find(self.target_key)
                if p1_index != -1:
                    insert_pos = p1_index + len(self.target_key)
                    new_name = new_name[:insert_pos] + prefix + new_name[insert_pos:]
                    has_changed = True
        
        return new_name + ext if has_changed else old_name, has_changed

    def _rename_file_directly(self, old_path):
        """直接重命名原文件（不复制、不保留副本）"""
        try:
            old_name = os.path.basename(old_path)
            ext = os.path.splitext(old_name)[1]
            
            # 仅处理docx文件
            if ext.lower() != ".docx":
                return False, f"跳过：非docx文件 - {old_name}"
            
            new_name, has_changed = self._get_new_filename(old_name)
            if not has_changed:
                return True, f"无需重命名：{old_name}"
            
            new_path = os.path.join(os.path.dirname(old_path), new_name)
            
            # 避免重名（直接覆盖已存在的文件）
            if os.path.exists(new_path):
                os.remove(new_path)
                self._log(f"  ⚠️  已删除同名文件：{new_name}")
            
            # 直接重命名原文件
            os.rename(old_path, new_path)
            return True, f"成功重命名：{old_name} → {new_name}"
        
        except Exception as e:
            return False, f"重命名失败：{old_name} - {str(e)}"

    def _batch_rename(self):
        """批量重命名文件（直接修改原文件）"""
        source_folder = self.source_folder.get().strip()
        
        if not source_folder or not os.path.isdir(source_folder):
            raise Exception("源文件夹无效，请选择有效的源文件夹")
        
        # 遍历源文件夹中的文件
        all_files = [f for f in os.listdir(source_folder) 
                     if os.path.isfile(os.path.join(source_folder, f))]
        
        if not all_files:
            raise Exception("源文件夹内未找到任何文件")
        
        # 开始批量处理
        success_count = 0
        skip_count = 0
        fail_count = 0
        
        self._log(f"\n📁 开始批量重命名（共{len(all_files)}个文件，直接修改原文件）")
        for filename in all_files:
            old_path = os.path.join(source_folder, filename)
            success, msg = self._rename_file_directly(old_path)
            self._log(f"  {msg}")
            
            if "成功" in msg or "无需重命名" in msg:
                success_count += 1
            elif "跳过" in msg:
                skip_count += 1
            elif "失败" in msg:
                fail_count += 1
        
        result = f"重命名完成 | 成功：{success_count} | 跳过：{skip_count} | 失败：{fail_count}"
        self._log(f"✅ {result}")
        return success_count > 0

    # -------------------------- 第二部分：内容处理功能（直接覆盖原文件） --------------------------
    def _remove_text(self, doc, text_to_remove):
        """删除文档中指定文本"""
        # 遍历所有段落
        for para in doc.paragraphs:
            if text_to_remove in para.text:
                para.text = para.text.replace(text_to_remove, "")
        # 遍历所有表格中的单元格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if text_to_remove in cell.text:
                        cell.text = cell.text.replace(text_to_remove, "")

    def _batch_replace_text(self, doc, replace_pairs):
        """批量替换文本"""
        # 替换段落中的文本
        for para in doc.paragraphs:
            for old_text, new_text in replace_pairs.items():
                if old_text in para.text:
                    para.text = para.text.replace(old_text, new_text)
        # 替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_text, new_text in replace_pairs.items():
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)

    def _remove_table_columns(self, doc, start_col, end_col):
        """删除表格中指定范围的列（索引从0开始）"""
        for table in doc.tables:
            max_cols = max(len(row.cells) for row in table.rows)
            if start_col >= max_cols:
                self._log(f"  ⚠️  表格列数不足，跳过列删除操作（当前最大列数: {max_cols}）")
                continue
            actual_end_col = min(end_col, max_cols - 1)
            # 从后往前删除列（避免索引错乱）
            for col_idx in range(actual_end_col, start_col - 1, -1):
                for row in table.rows:
                    if len(row.cells) > col_idx:
                        cell = row.cells[col_idx]
                        cell._element.getparent().remove(cell._element)

    def _swap_table_columns(self, doc, col1, col2):
        """交换表格中指定两列的内容"""
        for table in doc.tables:
            max_cols = max(len(row.cells) for row in table.rows)
            if col1 >= max_cols or col2 >= max_cols:
                self._log(f"  ⚠️  表格列数不足（当前最大列数: {max_cols}），跳过列交换操作")
                continue
            for row in table.rows:
                if len(row.cells) > max(col1, col2):
                    temp_text = row.cells[col1].text
                    row.cells[col1].text = row.cells[col2].text
                    row.cells[col2].text = temp_text

    def _execute_change_table_functions(self, doc):
        """执行基础文本/表格处理"""
        self._log("  📌 执行基础文本/表格处理")
        
        # 1. 删除所有"Test Report"文本
        self._remove_text(doc, "Test Report")
        self._log("  - 已删除所有'Test Report'文本")
        
        # 2. 删除第一个表格
        if doc.tables:
            first_table = doc.tables[0]
            table_element = first_table._element
            table_element.getparent().remove(table_element)
            self._log("  - 已删除第一个表格")
        else:
            self._log("  - 文档中未找到表格，跳过删除第一个表格操作")
            
        # 3. 批量替换文本
        replace_pairs = {
            "Final_Result": "",
            "Frequency": "频率",
            "QuasiPeak": "准峰值",
            "Margin": "裕量",
            "Limit": "限值"
        }
        self._batch_replace_text(doc, replace_pairs)
        self._log("  - 已完成文本批量替换")
        
        # 4. 删除所有表格的第5列到第9列（索引从0开始，对应4-8）
        self._remove_table_columns(doc, start_col=4, end_col=8)
        self._log("  - 已删除所有表格的第5列到第9列")
        
        # 5. 交换所有表格的第3列和第4列内容（索引2和3）
        self._swap_table_columns(doc, col1=2, col2=3)
        self._log("  - 已交换所有表格的第3列和第4列内容")

    def _set_cell_border(self, cell):
        """为单元格设置完整黑色边框（0.5磅实线）"""
        borders = ["top", "bottom", "left", "right"]
        for border_name in borders:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "000000")
            border.set(qn("w:space"), "0")
            cell._tc.get_or_add_tcPr().append(border)

    def _get_file_config(self, file_name):
        """根据文件名匹配配置项"""
        for keyword in self.table_config.keys():
            if keyword in file_name:
                return self.table_config[keyword]
        return self.default_config

    def _rebuild_table(self, table, table_data):
        """重建表格：原1-3列+新增3列+原4列（移至第7列）"""
        original_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            original_data.append(row_data)
        
        if not original_data:
            return None
        
        new_table_data = []
        new_col_headers = ["天线高度(cm)", "天线极化", "转台角度(deg)"]
        for idx, row in enumerate(original_data):
            # 补全原行至4列
            row += [""] * (4 - len(row))
            
            # 表头行填标题，数据行（1-6行）填指定值
            if idx == 0:
                new_cols = new_col_headers
            elif 1 <= idx <= 6:
                new_cols = table_data
            else:
                new_cols = ["", "", ""]
            
            # 新行结构：原1-3列 + 新增3列 + 原4列
            new_row = row[0:3] + new_cols + [row[3]]
            new_table_data.append(new_row)
        
        return new_table_data

    def _add_remark_row(self, table, remark_text):
        """添加第八行备注（合并所有列，文字靠左对齐）"""
        new_row = table.add_row().cells
        col_count = len(table.columns)
        
        # 合并所有列
        for i in range(1, col_count):
            new_row[0].merge(new_row[i])
        
        # 设置备注内容和格式
        cell = new_row[0]
        cell.text = remark_text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # 文字靠左对齐
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # 为合并单元格添加边框
        self._set_cell_border(cell)

    def _find_first_image(self, doc):
        """找到文档中第一个图片的XML节点和位置"""
        for idx, elem in enumerate(doc.element.body):
            if elem.tag.endswith("p"):
                for child in elem:
                    if child.tag.endswith("r"):
                        for grandchild in child:
                            if grandchild.tag.endswith("drawing"):
                                return elem, idx
        return None, -1

    def _process_table_and_image(self, doc, table, file_config):
        """核心：表格移至图片上方（间隔3行）+图片标注"""
        # 1. 找到第一个图片的位置
        img_para, img_idx = self._find_first_image(doc)
        if img_para is None:
            self._log("  ⚠️  未找到文档中的图片，跳过图片相关处理")
            doc.element.body.append(table._element)
            return
        
        self._log(f"  🖼️  找到图片，位置索引：{img_idx}")
        
        # 2. 先删除原表格（避免重复）
        old_table_elem = table._element
        if old_table_elem in doc.element.body:
            doc.element.body.remove(old_table_elem)
        
        # 3. 在图片上方插入3个空段落（间隔）
        for _ in range(3):
            empty_para = OxmlElement("w:p")
            doc.element.body.insert(img_idx, empty_para)
            img_idx += 1
        
        # 4. 将修改后的表格插入图片上方
        doc.element.body.insert(img_idx, table._element)
        
        # 5. 处理图片标注
        self._add_image_annotations(doc, img_para, file_config["polarization"])

    def _add_image_annotations(self, doc, img_para, polarization):
        """为图片添加标注：左上角“试验结果图：”+下方“水平/垂直极化”"""
        # 1. 图片左上角标注：试验结果图：（靠左）
        label_para = OxmlElement("w:p")
        label_run = OxmlElement("w:r")
        label_text = OxmlElement("w:t")
        label_text.text = "试验结果图："
        label_run.append(label_text)
        label_para.append(label_run)
        # 设置靠左对齐
        justify = OxmlElement("w:jc")
        justify.set(qn("w:val"), "left")
        label_para.append(justify)
        img_para.addprevious(label_para)
        
        # 2. 图片下方标注：水平/垂直极化（居中）
        polar_para = OxmlElement("w:p")
        polar_run = OxmlElement("w:r")
        polar_text = OxmlElement("w:t")
        polar_text.text = polarization
        polar_run.append(polar_text)
        polar_para.append(polar_run)
        # 设置居中对齐
        justify = OxmlElement("w:jc")
        justify.set(qn("w:val"), "center")
        polar_para.append(justify)
        img_para.addnext(polar_para)
        
        self._log(f"  ✅ 图片标注添加完成：试验结果图： + {polarization}")

    def _execute_2py_functions(self, doc, file_path, file_config):
        """执行表格扩展+图片处理"""
        self._log("  📌 执行表格扩展+图片处理")
        
        if not doc.tables:
            self._log("  ⚠️  文档中无表格，跳过表格扩展处理")
            return False
        
        # 处理第一个表格（核心处理对象）
        table = doc.tables[0]
        self._log(f"  📋 处理表格（原行列：{len(table.rows)}×{len(table.columns)}）")
        
        if len(table.columns) < 4:
            self._log(f"  ⚠️  表格列数不足4列，跳过表格扩展")
            return False
        
        # 重建表格数据
        new_table_data = self._rebuild_table(table, file_config["table_data"])
        if not new_table_data:
            self._log(f"  ⚠️  表格无数据，跳过表格扩展")
            return False
        
        # 删除原表格
        table_elem = table._element
        table_parent = table_elem.getparent()
        table_parent.remove(table_elem)
        
        # 创建新表格
        new_table = doc.add_table(rows=len(new_table_data), cols=7)
        new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置列宽
        for col in new_table.columns:
            col.width = Pt(60)
        
        # 填充数据+设置边框
        for row_idx, row_data in enumerate(new_table_data):
            row_cells = new_table.rows[row_idx].cells
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row_cells):
                    cell = row_cells[col_idx]
                    cell.text = cell_text
                    self._set_cell_border(cell)
        
        # 添加第八行备注
        self._add_remark_row(new_table, file_config["remark"])
        
        # 调整表格和图片位置+图片标注
        self._process_table_and_image(doc, new_table, file_config)
        
        return True

    def _modify_docx_content(self, file_path):
        """处理单个docx文件内容（直接覆盖原文件，不备份）"""
        # 1. 打开文档
        doc = Document(file_path)
        
        try:
            # 第一步：基础文本表格处理
            self._execute_change_table_functions(doc)
            
            # 第二步：表格扩展和图片处理
            file_config = self._get_file_config(os.path.basename(file_path))
            self._log(f"  📌 匹配关键词：{[k for k in self.table_config if k in file_path] or '无'}")
            self._execute_2py_functions(doc, file_path, file_config)
            
            # 直接保存覆盖原文件
            doc.save(file_path)
            self._log(f"  ✅ 内容处理完成（已覆盖原文件）：{os.path.basename(file_path)}")
            return True
        except Exception as e:
            self._log(f"❌ 内容处理失败：{str(e)}")
            self._log(f"❌ 错误详情：{traceback.format_exc()}")
            return False

    def _batch_process_content(self, folder):
        """批量处理文件夹下docx文件的内容（直接覆盖原文件）"""
        # 筛选docx文件
        docx_files = [
            f for f in os.listdir(folder)
            if f.lower().endswith(".docx") and os.path.isfile(os.path.join(folder, f))
        ]
        
        if not docx_files:
            raise Exception("源文件夹中未找到任何docx文件")
        
        self._log(f"\n📊 开始内容处理（共{len(docx_files)}个docx文件，直接覆盖原文件）")
        
        # 批量处理
        success = 0
        fail = 0
        for file_name in docx_files:
            file_path = os.path.join(folder, file_name)
            self._log(f"\n🔍 处理文件：{file_name}")
            try:
                if self._modify_docx_content(file_path):
                    success += 1
                else:
                    fail += 1
            except Exception as e:
                self._log(f"❌ 处理异常：{str(e)}")
                fail += 1
        
        result = f"内容处理完成 | 成功：{success} | 失败：{fail}"
        self._log(f"✅ {result}")
        return success, fail

    # -------------------------- 主处理逻辑 --------------------------
    def _batch_process_all(self):
        """完整处理流程：先重命名，再处理内容（均直接修改原文件）"""
        # 二次确认：防止误操作
        confirm = messagebox.askyesno(
            "危险操作确认", 
            "此操作会直接修改源文件夹中的所有docx文件，且不会保留原文件备份！\n请确认已做好数据备份，是否继续？"
        )
        if not confirm:
            self._log("📌 用户取消了操作")
            return
        
        try:
            # 清空日志
            self._clear_log()
            self._log("🚀 开始Docx批量处理（直接修改原文件，无备份）")
            
            # 第一步：批量重命名（直接修改原文件）
            rename_success = self._batch_rename()
            
            if not rename_success:
                self._log("⚠️  重命名无成功文件，跳过内容处理")
                messagebox.showwarning("警告", "重命名无成功文件，跳过内容处理")
                return
            
            # 第二步：批量处理内容（直接覆盖原文件）
            content_success, content_fail = self._batch_process_content(self.source_folder.get().strip())
            
            # 最终结果
            total_result = f"""
处理完成！
📝 重命名：成功（见日志）
📝 内容修改：成功 {content_success} 个 | 失败 {content_fail} 个
⚠️  所有修改均直接覆盖原文件，未保留备份！
"""
            self._log(f"\n{total_result}")
            messagebox.showinfo("处理完成", total_result)
            
        except Exception as e:
            self._log(f"❌ 整体处理失败：{str(e)}")
            messagebox.showerror("错误", f"处理失败：{str(e)}")

if __name__ == "__main__":
    # 适配tkinter中文显示
    root = tk.Tk()
    root.option_add("*Font", "SimHei 9")
    app = DocxBatchProcessor(root)
    root.mainloop()
