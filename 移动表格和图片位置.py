#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Docx批量处理工具 - 表格图片调整 + 清除页眉页脚
Python 3.8.7 + python-docx 0.8.11
核心功能：
1. 批量处理文件夹内所有docx文件
2. 清除第一个表格上方所有空段落，确保表格距离顶部有2个空段落
3. 表格移至图片上方（无间隔行）
4. 图片左上角标注：试验结果图：
5. 图片下方中间标注：水平极化
6. “试验结果图：”前面保留一个空行
7. 新增：删除页眉和页尾的所有内容
"""
import os
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class DocxBatchTool:
    def __init__(self, root):
        self.root = root
        self.root.title("Docx表格图片批量处理工具")
        self.root.geometry("800x650")
        
        # 配置项（移除了间隔行数配置）
        self.img_label_top = "\n试验结果图："       # 图片左上角文字
        self.img_label_bottom = "水平极化"        # 图片下方中间文字
        self.table_top_spaces = 2                # 表格距离顶部的空段落数
        
        self.folder_path = tk.StringVar()
        self._build_gui()

    def _build_gui(self):
        """构建GUI界面"""
        # 1. 文件夹选择区域
        frame1 = tk.Frame(self.root, padx=10, pady=10)
        frame1.pack(fill=tk.X)
        
        tk.Label(frame1, text="目标文件夹：", font=("SimHei", 10)).pack(side=tk.LEFT)
        tk.Entry(frame1, textvariable=self.folder_path, width=65, font=("SimHei", 10)).pack(side=tk.LEFT, padx=5)
        tk.Button(
            frame1, text="选择文件夹", command=self._select_folder,
            font=("SimHei", 10), bg="#E0E0E0"
        ).pack(side=tk.LEFT)
        
        # 2. 执行按钮
        frame2 = tk.Frame(self.root, padx=10, pady=8)
        frame2.pack(fill=tk.X)
        
        tk.Button(
            frame2, text="开始批量处理", 
            command=self._batch_process,
            bg="#2196F3", fg="white", font=("SimHei", 11, "bold"), padx=30
        ).pack(side=tk.LEFT)
        
        # 3. 日志显示区域
        frame3 = tk.Frame(self.root, padx=10, pady=10)
        frame3.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(frame3, text="处理日志：", font=("SimHei", 10)).pack(anchor=tk.W)
        self.log_text = scrolledtext.ScrolledText(
            frame3, height=30, font=("Consolas", 9), wrap=tk.WORD
        )
        self.log_text.pack(fill=tk.BOTH, expand=True)

    def _select_folder(self):
        """选择目标文件夹"""
        folder = filedialog.askdirectory(title="选择包含docx文件的文件夹")
        if folder:
            self.folder_path.set(folder)
            self._log(f"✅ 已选择文件夹：{folder}")

    def _log(self, msg):
        """日志输出（自动滚动）"""
        self.log_text.insert(tk.END, f"{msg}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
    
    def _clear_header_footer(self, doc):
        """删除页眉和页脚的所有内容"""
        try:
            # 处理页眉
            header_removed = 0
            for section in doc.sections:
                header = section.header
                # 清空页眉所有段落
                for para in header.paragraphs:
                    para.clear()
                    header_removed += 1
                # 处理页眉中的表格（如果有）
                for table in header.tables:
                    header._element.remove(table._element)
                    header_removed += 1
            
            # 处理页脚
            footer_removed = 0
            for section in doc.sections:
                footer = section.footer
                # 清空页脚所有段落
                for para in footer.paragraphs:
                    para.clear()
                    footer_removed += 1
                # 处理页脚中的表格（如果有）
                for table in footer.tables:
                    footer._element.remove(table._element)
                    footer_removed += 1
            
            self._log(f"  ✅ 清除页眉内容数：{header_removed} | 清除页脚内容数：{footer_removed}")
            return True
        except Exception as e:
            self._log(f"  ⚠️  清除页眉页脚失败：{str(e)}")
            return False
    
    def _find_first_image(self, doc):
        """精准定位文档中第一个图片的段落（支持所有图片格式）"""
        self._log("  🔍 开始定位图片...")
        
        # 方法1：遍历段落+run（主流嵌入式图片）
        for para_idx, para in enumerate(doc.paragraphs):
            for run in para.runs:
                # 检测drawing（2007+）和pict（老式）图片
                if run.element.xpath(".//w:drawing") or run.element.xpath(".//w:pict"):
                    self._log(f"    ✅ 在段落 {para_idx+1} 找到图片")
                    return para
        
        # 方法2：直接遍历XML（兜底方案）
        self._log("    ⚠️  Run中未找到图片，尝试遍历文档XML...")
        body = doc.element.body
        for elem in body.iter():
            if elem.tag.endswith('drawing') or elem.tag.endswith('pict'):
                # 向上查找包含图片的段落
                para_elem = elem.getparent()
                while para_elem is not None and not para_elem.tag.endswith('p'):
                    para_elem = para_elem.getparent()
                if para_elem is not None:
                    # 转换为Paragraph对象
                    for para in doc.paragraphs:
                        if para._p == para_elem:
                            self._log(f"    ✅ 在XML中找到图片，对应段落")
                            return para
        
        self._log("    ❌ 未找到任何图片！")
        return None

    def _add_image_annotations(self, doc, img_para):
        """为图片添加标注：左上角+下方中间"""
        try:
            # 1. 图片左上角标注（试验结果图：）- 靠左对齐
            top_para = doc.add_paragraph()
            top_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            top_run = top_para.add_run(self.img_label_top)
            top_run.font.size = Pt(10)
            top_run.font.name = "宋体"
            # 插入到图片段落正上方
            img_para._p.addprevious(top_para._p)
            
            # 2. 图片下方中间标注（水平极化）- 居中对齐
            bottom_para = doc.add_paragraph()
            bottom_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            bottom_run = bottom_para.add_run(self.img_label_bottom)
            bottom_run.font.size = Pt(10)
            bottom_run.font.name = "宋体"
            # 插入到图片段落正下方
            img_para._p.addnext(bottom_para._p)
            
            self._log(f"  ✅ 图片标注完成：{self.img_label_top} + {self.img_label_bottom}")
            return True
        except Exception as e:
            self._log(f"  ⚠️  图片标注失败：{str(e)}")
            return False

    def _clear_empty_paragraphs_above_table(self, doc, table):
        """清除第一个表格上方的所有空段落，并确保表格距离文档顶部有2个空段落"""
        try:
            # 获取表格对应的XML元素
            table_elem = table._element
            parent_elem = table_elem.getparent()
            # 找到表格在父元素中的索引
            table_index = list(parent_elem).index(table_elem)
            
            # 步骤1：从表格上方开始向前遍历，清理所有空段落
            removed_count = 0
            for i in range(table_index - 1, -1, -1):
                elem = parent_elem[i]
                # 判断是否是空段落（无有效内容）
                if elem.tag.endswith('p'):
                    # 检查段落是否为空（无文字/仅空白符）
                    para_text = ""
                    for run in elem.xpath(".//w:t"):
                        para_text += run.text or ""
                    if not para_text.strip():
                        # 删除空段落
                        parent_elem.remove(elem)
                        removed_count += 1
                        # 移除后表格索引会变化，需要重新计算
                        table_index = list(parent_elem).index(table_elem)
            
            if removed_count > 0:
                self._log(f"  ✅ 清除表格上方空行数量：{removed_count}")
            else:
                self._log(f"  ℹ️  表格上方无空行需要清除")
            
            # 步骤2：确保表格顶部有且仅有2个空段落
            # 重新获取清理后的表格索引
            table_index = list(parent_elem).index(table_elem)
            # 统计表格上方已有的非空段落数量（向上遍历直到文档顶部）
            non_empty_above = 0
            for i in range(table_index - 1, -1, -1):
                elem = parent_elem[i]
                if elem.tag.endswith('p'):
                    # 检查是否为非空段落
                    para_text = ""
                    for run in elem.xpath(".//w:t"):
                        para_text += run.text or ""
                    if para_text.strip():
                        non_empty_above += 1
            
            # 计算需要插入的空段落数量（目标：表格上方有2个空段落）
            current_empty = table_index - non_empty_above
            insert_count = self.table_top_spaces - current_empty
            
            if insert_count > 0:
                # 倒序插入空段落（保证顺序正确）
                for i in reversed(range(insert_count)):
                    empty_para = parse_xml(f'<w:p {nsdecls("w")}/>')
                    parent_elem.insert(table_index, empty_para)
                self._log(f"  ✅ 插入{insert_count}个空段落，确保表格顶部有{self.table_top_spaces}个空行")
            elif insert_count < 0:
                # 理论上不会触发（已清理所有空段落），仅做兜底
                self._log(f"  ℹ️  表格上方空行已超过{self.table_top_spaces}个，无需调整")
            
            return True
        except Exception as e:
            self._log(f"  ⚠️  清除/调整表格上方空行失败：{str(e)}")
            return False

    def _process_single_file(self, file_path):
        """处理单个docx文件"""
        try:
            file_name = os.path.basename(file_path)
            self._log(f"\n===== 处理文件：{file_name} =====")
            
            # 1. 备份原文件（防止数据丢失）
            backup_path = f"{file_path}.bak"
            shutil.copy2(file_path, backup_path)
            self._log(f"  📁 已备份原文件：{file_name}.bak")
            
            # 2. 打开文档
            doc = Document(file_path)
            self._log(f"  📄 文档段落数：{len(doc.paragraphs)} | 表格数：{len(doc.tables)}")
            
            # 3. 新增功能：删除页眉页脚所有内容
            self._clear_header_footer(doc)
            
            # 4. 定位图片
            img_para = self._find_first_image(doc)
            
            # 5. 核心：表格移至图片上方（无间隔行）+ 图片标注
            table = doc.tables[0] if doc.tables else None
            if table and img_para:
                # 移除原表格，直接插入到图片上方（无间隔行）
                table_elem = table._element
                table_elem.getparent().remove(table_elem)
                img_para._p.addprevious(table_elem)
                self._log(f"  ✅ 表格已移至图片上方（无间隔行）")
                # 添加图片标注
                self._add_image_annotations(doc, img_para)
            elif not table:
                self._log("  ⚠️  无表格，仅处理图片标注")
                if img_para:
                    self._add_image_annotations(doc, img_para)
            elif not img_para:
                self._log("  ⚠️  未找到图片，跳过表格移动和标注")

            # 6. 处理表格：清除表格上方空行并保证顶部2个空段落
            table = doc.tables[0] if doc.tables else None
            if table:
                self._clear_empty_paragraphs_above_table(doc, table)
            else:
                self._log("  ⚠️  文档中无表格，跳过空行清理")
            
            # 7. 保存修改后的文档
            doc.save(file_path)
            self._log(f"  ✅ 文件处理完成：{file_name}")
            return True
        except Exception as e:
            self._log(f"❌ 文件处理异常：{str(e)}")
            import traceback
            self._log(f"📝 详细错误：{traceback.format_exc()[:500]}")  # 限制错误日志长度
            return False

    def _batch_process(self):
        """批量处理文件夹下所有docx文件"""
        folder = self.folder_path.get()
        if not folder or not os.path.isdir(folder):
            messagebox.showerror("错误", "请选择有效的文件夹！")
            return
        
        # 清空日志
        self.log_text.delete(1.0, tk.END)
        self._log("🚀 开始批量处理docx文件...")
        self._log(f"📂 目标文件夹：{folder}")
        
        # 筛选所有docx文件
        docx_files = [
            os.path.join(folder, f) for f in os.listdir(folder)
            if f.lower().endswith(".docx") and os.path.isfile(os.path.join(folder, f))
        ]
        
        if not docx_files:
            self._log("⚠️  未找到任何.docx文件！")
            messagebox.showinfo("提示", "未找到任何.docx文件！")
            return
        
        self._log(f"📊 共找到 {len(docx_files)} 个docx文件")
        
        # 批量处理
        success_count = 0
        fail_count = 0
        for file_path in docx_files:
            if self._process_single_file(file_path):
                success_count += 1
            else:
                fail_count += 1
        
        # 处理完成统计
        result_msg = f"\n✅ 批量处理完成！成功：{success_count}个 | 失败：{fail_count}个"
        self._log(result_msg)
        messagebox.showinfo("处理完成", result_msg)

if __name__ == "__main__":
    # 适配tkinter中文显示
    root = tk.Tk()
    root.option_add("*Font", "SimHei 9")
    # 启动主程序
    app = DocxBatchTool(root)
    root.mainloop()
