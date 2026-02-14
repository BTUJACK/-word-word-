#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Docx批量处理工具 - 表格图片调整
Python 3.8.7 + python-docx 0.8.11
核心功能：
1. 批量处理文件夹内所有docx文件
2. 强制显示表格所有边框（黑色0.5磅实线）
3. 表格移至图片上方，与图片间隔3行
4. 图片左上角标注：试验结果图：
5. 图片下方中间标注：水平极化
"""
import os
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.shared import OxmlElement, qn

class DocxBatchTool:
    def __init__(self, root):
        self.root = root
        self.root.title("Docx表格图片批量处理工具")
        self.root.geometry("800x650")
        
        # 配置项
        self.img_label_top = "试验结果图："       # 图片左上角文字
        self.img_label_bottom = "水平极化"        # 图片下方中间文字
        self.space_lines = 3                     # 表格与图片的间隔行数
        
        self.folder_path = tk.StringVar()
        self._build_gui()

    def _build_gui(self):
        """构建GUI界面"""
        # 1. 文件夹选择区域
        frame1 = tk.Frame(self.root, padx=10, pady=10)
        frame1.pack(fill=tk.X)
        
        tk.Label(frame1, text="目标文件夹：", font=("SimHei", 10)).pack(side=tk.LEFT)
        tk.Entry(frame1, textvariable=self.folder_path, width=65, font=("SimHei", 10)).pack(side=tk.LEFT, padx=5)
        tk.Button(
            frame1, text="选择文件夹", command=self._select_folder,
            font=("SimHei", 10), bg="#E0E0E0"
        ).pack(side=tk.LEFT)
        
        # 2. 执行按钮
        frame2 = tk.Frame(self.root, padx=10, pady=8)
        frame2.pack(fill=tk.X)
        
        tk.Button(
            frame2, text="开始批量处理", 
            command=self._batch_process,
            bg="#2196F3", fg="white", font=("SimHei", 11, "bold"), padx=30
        ).pack(side=tk.LEFT)
        
        # 3. 日志显示区域
        frame3 = tk.Frame(self.root, padx=10, pady=10)
        frame3.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(frame3, text="处理日志：", font=("SimHei", 10)).pack(anchor=tk.W)
        self.log_text = scrolledtext.ScrolledText(
            frame3, height=30, font=("Consolas", 9), wrap=tk.WORD
        )
        self.log_text.pack(fill=tk.BOTH, expand=True)

    def _select_folder(self):
        """选择目标文件夹"""
        folder = filedialog.askdirectory(title="选择包含docx文件的文件夹")
        if folder:
            self.folder_path.set(folder)
            self._log(f"✅ 已选择文件夹：{folder}")

    def _log(self, msg):
        """日志输出（自动滚动）"""
        self.log_text.insert(tk.END, f"{msg}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def _set_cell_border(self, cell):
        """为单元格设置完整黑色边框（0.5磅实线）"""
        try:
            tcPr = cell._tc.get_or_add_tcPr()
            
            # 清除原有边框（避免样式冲突）
            for border in tcPr.findall(".//*[local-name()='top' or local-name()='bottom' or local-name()='left' or local-name()='right']"):
                tcPr.remove(border)
            
            # 边框样式：黑色、0.5磅、实线
            border_style = {
                "val": "single",
                "sz": "4",       # 0.5磅（1pt=8sz）
                "color": "000000",  # 黑色
                "space": "0"
            }
            
            # 为四个方向添加边框
            for border_name in ["top", "bottom", "left", "right"]:
                border = OxmlElement(f"w:{border_name}")
                for key, value in border_style.items():
                    border.set(qn(f"w:{key}"), value)
                tcPr.append(border)
            
            # 单元格文字垂直居中
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        except Exception as e:
            self._log(f"  ⚠️  单元格边框设置失败：{str(e)}")

    def _apply_table_borders(self, table):
        """为整个表格的所有单元格添加边框"""
        try:
            for row in table.rows:
                for cell in row.cells:
                    self._set_cell_border(cell)
            self._log("  ✅ 表格边框已全部显示（黑色0.5磅实线）")
        except Exception as e:
            self._log(f"  ⚠️  表格边框设置失败：{str(e)}")

    def _find_first_image(self, doc):
        """精准定位文档中第一个图片的段落（支持所有图片格式）"""
        self._log("  🔍 开始定位图片...")
        
        # 方法1：遍历段落+run（主流嵌入式图片）
        for para_idx, para in enumerate(doc.paragraphs):
            for run in para.runs:
                # 检测drawing（2007+）和pict（老式）图片
                if run.element.xpath(".//w:drawing") or run.element.xpath(".//w:pict"):
                    self._log(f"    ✅ 在段落 {para_idx+1} 找到图片")
                    return para
        
        # 方法2：直接遍历XML（兜底方案）
        self._log("    ⚠️  Run中未找到图片，尝试遍历文档XML...")
        body = doc.element.body
        for elem in body.iter():
            if elem.tag.endswith('drawing') or elem.tag.endswith('pict'):
                # 向上查找包含图片的段落
                para_elem = elem.getparent()
                while para_elem is not None and not para_elem.tag.endswith('p'):
                    para_elem = para_elem.getparent()
                if para_elem is not None:
                    # 转换为Paragraph对象
                    for para in doc.paragraphs:
                        if para._p == para_elem:
                            self._log(f"    ✅ 在XML中找到图片，对应段落")
                            return para
        
        self._log("    ❌ 未找到任何图片！")
        return None

    def _insert_space_paragraphs(self, doc, ref_para, count):
        """在参考段落上方插入指定数量的空段落（间隔）"""
        try:
            ref_elem = ref_para._p
            parent_elem = ref_elem.getparent()
            ref_index = list(parent_elem).index(ref_elem)
            
            # 倒序插入空段落（保证顺序正确）
            for i in reversed(range(count)):
                empty_para = parse_xml(f'<w:p {nsdecls("w")}/>')
                parent_elem.insert(ref_index, empty_para)
            
            self._log(f"  ✅ 插入{count}个空段落（表格与图片间隔）")
            return ref_index
        except Exception as e:
            self._log(f"  ⚠️  空段落插入失败：{str(e)}")
            return -1

    def _add_image_annotations(self, doc, img_para):
        """为图片添加标注：左上角+下方中间"""
        try:
            # 1. 图片左上角标注（试验结果图：）- 靠左对齐
            top_para = doc.add_paragraph()
            top_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            top_run = top_para.add_run(self.img_label_top)
            top_run.font.size = Pt(10)
            top_run.font.name = "宋体"
            # 插入到图片段落正上方
            img_para._p.addprevious(top_para._p)
            
            # 2. 图片下方中间标注（水平极化）- 居中对齐
            bottom_para = doc.add_paragraph()
            bottom_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            bottom_run = bottom_para.add_run(self.img_label_bottom)
            bottom_run.font.size = Pt(10)
            bottom_run.font.name = "宋体"
            # 插入到图片段落正下方
            img_para._p.addnext(bottom_para._p)
            
            self._log(f"  ✅ 图片标注完成：{self.img_label_top} + {self.img_label_bottom}")
            return True
        except Exception as e:
            self._log(f"  ⚠️  图片标注失败：{str(e)}")
            return False

    def _process_single_file(self, file_path):
        """处理单个docx文件"""
        try:
            file_name = os.path.basename(file_path)
            self._log(f"\n===== 处理文件：{file_name} =====")
            
            # 1. 备份原文件（防止数据丢失）
            backup_path = f"{file_path}.bak"
            shutil.copy2(file_path, backup_path)
            self._log(f"  📁 已备份原文件：{file_name}.bak")
            
            # 2. 打开文档
            doc = Document(file_path)
            self._log(f"  📄 文档段落数：{len(doc.paragraphs)} | 表格数：{len(doc.tables)}")
            
            # 3. 处理表格边框
            table = doc.tables[0] if doc.tables else None
            if table:
                self._apply_table_borders(table)
            else:
                self._log("  ⚠️  文档中无表格，跳过边框设置")
            
            # 4. 定位图片
            img_para = self._find_first_image(doc)
            
            # 5. 核心：表格移至图片上方（间隔3行）+ 图片标注
            if table and img_para:
                # 先插入间隔空段落
                self._insert_space_paragraphs(doc, img_para, self.space_lines)
                # 移除原表格，插入到图片上方（空段落之后）
                table_elem = table._element
                table_elem.getparent().remove(table_elem)
                img_para._p.addprevious(table_elem)
                self._log(f"  ✅ 表格已移至图片上方（间隔{self.space_lines}行）")
                # 添加图片标注
                self._add_image_annotations(doc, img_para)
            elif not table:
                self._log("  ⚠️  无表格，仅处理图片标注")
                if img_para:
                    self._add_image_annotations(doc, img_para)
            elif not img_para:
                self._log("  ⚠️  未找到图片，仅保留表格边框")
            
            # 6. 保存修改后的文档
            doc.save(file_path)
            self._log(f"  ✅ 文件处理完成：{file_name}")
            return True
        except Exception as e:
            self._log(f"❌ 文件处理异常：{str(e)}")
            import traceback
            self._log(f"📝 详细错误：{traceback.format_exc()[:500]}")  # 限制错误日志长度
            return False

    def _batch_process(self):
        """批量处理文件夹下所有docx文件"""
        folder = self.folder_path.get()
        if not folder or not os.path.isdir(folder):
            messagebox.showerror("错误", "请选择有效的文件夹！")
            return
        
        # 清空日志
        self.log_text.delete(1.0, tk.END)
        self._log("🚀 开始批量处理docx文件...")
        self._log(f"📂 目标文件夹：{folder}")
        
        # 筛选所有docx文件
        docx_files = [
            os.path.join(folder, f) for f in os.listdir(folder)
            if f.lower().endswith(".docx") and os.path.isfile(os.path.join(folder, f))
        ]
        
        if not docx_files:
            self._log("⚠️  未找到任何.docx文件！")
            messagebox.showinfo("提示", "未找到任何.docx文件！")
            return
        
        self._log(f"📊 共找到 {len(docx_files)} 个docx文件")
        
        # 批量处理
        success_count = 0
        fail_count = 0
        for file_path in docx_files:
            if self._process_single_file(file_path):
                success_count += 1
            else:
                fail_count += 1
        
        # 处理完成统计
        result_msg = f"\n✅ 批量处理完成！成功：{success_count}个 | 失败：{fail_count}个"
        self._log(result_msg)
        messagebox.showinfo("处理完成", result_msg)

if __name__ == "__main__":
    # 适配tkinter中文显示
    root = tk.Tk()
    root.option_add("*Font", "SimHei 9")
    # 启动主程序
    app = DocxBatchTool(root)
    root.mainloop()
