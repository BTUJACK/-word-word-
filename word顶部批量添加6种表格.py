#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Docx批量处理工具 - 保留图片+第二行插入表格
Python 3.8.7 + python-docx 0.8.11 完全兼容版
核心功能：
1. 匹配12种文件名关键词（Ambient/M1-M5 + ME_H/RE_H）
2. 第二行插入2列2行表格（第二行合并）+ 对应内容
3. 表格后保留2行空白内容
4. 完全保留文档原有图片、格式




用Python 3.8.7实现批量修改一个文件夹里面的.docx文件，并且生成一个GUI界面进行操作。
1
如果.docx文件含有“Ambient_ME _H”：
第二行添加两列两行的表格，并且第二行表格合并，并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：150kHz-30MHz
样品运行模式：背景噪声
    
表格后面保留两行空白内容；

2
如果.docx文件含有“Ambient_RE _H”：
第二行添加两列两行的表格，并且第二行表格合并，并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：30MHz-1GHz
样品运行模式：背景噪声

表格后面保留两行空白内容；

3
如果.docx文件含有“M1_ME_H”：
第二行添加两列两行的表格；第二行表格合并；并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：150kHz-30MHz
样品运行模式：1

表格后面保留两行空白内容；
4
如果.docx文件含有“M1_RE_H”：
第二行添加两列两行的表格，并且第二行表格合并，并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：30MHz-1GHz
样品运行模式：1

表格后面保留两行空白内容；

5
如果.docx文件含有“M2_ME_H”：
第二行添加两列两行的表格；第二行表格合并；并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：150kHz-30MHz
样品运行模式：2

表格后面保留两行空白内容；
6
如果.docx文件含有“M2_RE_H”：
第二行添加两列两行的表格，并且第二行表格合并，并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：30MHz-1GHz
样品运行模式：2

表格后面保留两行空白内容；

7
如果.docx文件含有“M3_ME_H”：
第二行添加两列两行的表格；第二行表格合并；并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：150kHz-30MHz
样品运行模式：3

表格后面保留两行空白内容；
8
如果.docx文件含有“M3_RE_H”：
第二行添加两列两行的表格，并且第二行表格合并，并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：30MHz-1GHz
样品运行模式：3

表格后面保留两行空白内容；

9
如果.docx文件含有“M4_ME_H”：
第二行添加两列两行的表格；第二行表格合并；并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：150kHz-30MHz
样品运行模式：4

表格后面保留两行空白内容；
10
如果.docx文件含有“M4_RE_H”：
第二行添加两列两行的表格，并且第二行表格合并，并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：30MHz-1GHz
样品运行模式：4

表格后面保留两行空白内容；

11
如果.docx文件含有“M5_ME_H”：
第二行添加两列两行的表格；第二行表格合并；并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：150kHz-30MHz
样品运行模式：5

表格后面保留两行空白内容；
12
如果.docx文件含有“M5_RE_H”：
第二行添加两列两行的表格，并且第二行表格合并，并且添加内容：
试验供电电源：380V AC/50Hz 试验频率范围：30MHz-1GHz
样品运行模式：5

表格后面保留两行空白内容；



"""
import os
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.shared import OxmlElement, qn

class DocxBatchTableTool:
    def __init__(self, root):
        self.root = root
        self.root.title("Docx批量添加表格工具（保留图片+第二行插入）")
        self.root.geometry("800x650")
        
        # 核心配置：12种关键词对应的表格内容（按优先级排序）
        self.keyword_content_map = {
            # Ambient系列
            "Ambient_ME_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：150kHz-30MHz",
                "row2_merged": "样品运行模式：背景噪声"
            },
            "Ambient_RE_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：30MHz-1GHz",
                "row2_merged": "样品运行模式：背景噪声"
            },
            # M1系列
            "M1_ME_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：150kHz-30MHz",
                "row2_merged": "样品运行模式：1"
            },
            "M1_RE_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：30MHz-1GHz",
                "row2_merged": "样品运行模式：1"
            },
            # M2系列
            "M2_ME_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：150kHz-30MHz",
                "row2_merged": "样品运行模式：2"
            },
            "M2_RE_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：30MHz-1GHz",
                "row2_merged": "样品运行模式：2"
            },
            # M3系列
            "M3_ME_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：150kHz-30MHz",
                "row2_merged": "样品运行模式：3"
            },
            "M3_RE_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：30MHz-1GHz",
                "row2_merged": "样品运行模式：3"
            },
            # M4系列
            "M4_ME_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：150kHz-30MHz",
                "row2_merged": "样品运行模式：4"
            },
            "M4_RE_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：30MHz-1GHz",
                "row2_merged": "样品运行模式：4"
            },
            # M5系列
            "M5_ME_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：150kHz-30MHz",
                "row2_merged": "样品运行模式：5"
            },
            "M5_RE_H": {
                "row1_col1": "试验供电电源：380V AC/50Hz",
                "row1_col2": "试验频率范围：30MHz-1GHz",
                "row2_merged": "样品运行模式：5"
            }
        }
        self.blank_lines_after_table = 2  # 表格后保留的空白行数
        
        self.folder_path = tk.StringVar()
        self._build_gui()

    def _build_gui(self):
        """构建GUI界面"""
        # 1. 文件夹选择区域
        frame1 = tk.Frame(self.root, padx=10, pady=10)
        frame1.pack(fill=tk.X)
        
        tk.Label(frame1, text="目标文件夹：", font=("SimHei", 10)).pack(side=tk.LEFT)
        tk.Entry(frame1, textvariable=self.folder_path, width=65, font=("SimHei", 10)).pack(side=tk.LEFT, padx=5)
        tk.Button(
            frame1, text="选择文件夹", command=self._select_folder,
            font=("SimHei", 10), bg="#E0E0E0"
        ).pack(side=tk.LEFT)
        
        # 2. 执行按钮
        frame2 = tk.Frame(self.root, padx=10, pady=8)
        frame2.pack(fill=tk.X)
        
        tk.Button(
            frame2, text="开始批量处理", 
            command=self._batch_process,
            bg="#2196F3", fg="white", font=("SimHei", 11, "bold"), padx=30
        ).pack(side=tk.LEFT)
        
        # 3. 日志显示区域
        frame3 = tk.Frame(self.root, padx=10, pady=10)
        frame3.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(frame3, text="处理日志：", font=("SimHei", 10)).pack(anchor=tk.W)
        self.log_text = scrolledtext.ScrolledText(
            frame3, height=30, font=("Consolas", 9), wrap=tk.WORD
        )
        self.log_text.pack(fill=tk.BOTH, expand=True)

    def _select_folder(self):
        """选择目标文件夹"""
        folder = filedialog.askdirectory(title="选择包含docx文件的文件夹")
        if folder:
            self.folder_path.set(folder)
            self._log(f"✅ 已选择文件夹：{folder}")

    def _log(self, msg):
        """日志输出（自动滚动）"""
        self.log_text.insert(tk.END, f"{msg}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def _check_filename_keyword(self, file_path):
        """检测文件名是否包含12种关键词（大小写不敏感）"""
        file_name = os.path.basename(file_path).lower()
        # 按配置顺序匹配，确保优先级
        for keyword in self.keyword_content_map.keys():
            if keyword.lower() in file_name:
                return keyword
        return None

    def _set_cell_border(self, cell):
        """手动为单元格添加黑色边框（不依赖预设样式）"""
        try:
            tcPr = cell._tc.get_or_add_tcPr()
            
            # 边框样式：黑色、0.5磅实线（兼容所有Word版本）
            border_style = {
                "val": "single",
                "sz": "4",       # 0.5磅（1pt=8sz）
                "color": "000000",  # 黑色
                "space": "0"
            }
            
            # 为单元格的四个方向添加边框
            for border_name in ["top", "bottom", "left", "right"]:
                border = OxmlElement(f"w:{border_name}")
                for key, value in border_style.items():
                    border.set(qn(f"w:{key}"), value)
                tcPr.append(border)
            
            # 单元格文字垂直居中
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        except Exception as e:
            self._log(f"  ⚠️  单元格边框设置失败：{str(e)}")

    def _apply_table_borders(self, table):
        """为整个表格的所有单元格添加边框"""
        try:
            for row in table.rows:
                for cell in row.cells:
                    self._set_cell_border(cell)
            self._log("  ✅ 表格边框已手动添加（黑色0.5磅实线）")
        except Exception as e:
            self._log(f"  ⚠️  表格边框设置失败：{str(e)}")

    def _insert_table_at_second_line(self, doc, keyword):
        """
        安全插入表格到第二行（保留图片）
        核心逻辑：先在文档末尾创建表格，再通过段落移动到第二行，避免破坏XML结构
        """
        try:
            # 1. 确保文档至少有1个段落（为第二行预留位置）
            if len(doc.paragraphs) == 0:
                doc.add_paragraph("")  # 第一行空段落占位
                self._log("  ⚠️  文档为空，先插入第一行空段落占位")
            
            # 2. 获取当前关键词对应的表格内容
            content = self.keyword_content_map[keyword]
            
            # 3. 先在文档末尾创建表格（避免破坏现有结构）
            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT  # 表格左对齐
            
            # 设置表格列宽（优化显示效果）
            for row in table.rows:
                row.cells[0].width = Inches(3.0)
                row.cells[1].width = Inches(3.0)
            
            # 4. 合并第二行的两个单元格
            row2_cells = table.rows[1].cells
            row2_cells[0].merge(row2_cells[1])
            
            # 5. 手动添加表格边框
            self._apply_table_borders(table)
            
            # 6. 填充表格内容
            # 第一行第一列
            cell1_1 = table.cell(0, 0)
            cell1_1.text = content["row1_col1"]
            # 第一行第二列
            cell1_2 = table.cell(0, 1)
            cell1_2.text = content["row1_col2"]
            # 第二行（合并后）
            cell2 = table.cell(1, 0)  # 合并后仅需操作第一个单元格
            cell2.text = content["row2_merged"]
            
            # 统一设置表格文字样式（宋体10号）
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = "宋体"
                            run.font.size = Pt(10)
            
            # 7. 安全移动表格到第二行（核心修复：保留图片）
            # 获取表格的XML元素
            table_elem = table._tbl
            # 从末尾移除表格
            table_elem.getparent().remove(table_elem)
            
            # 获取文档body的子元素列表（包含段落、图片、表格等所有元素）
            body_children = list(doc._body._element)
            
            # 找到第二行的位置插入表格（索引1）
            if len(body_children) >= 1:
                # 在第一个元素（第一行）之后插入表格（第二行）
                doc._body._element.insert(1, table_elem)
            else:
                # 若只有1个元素，插入到第一个元素之后
                doc._body._element.append(table_elem)
            
            self._log("  ✅ 表格已安全移动到第二行（保留图片）")
            
            # 8. 表格后插入2行空白（紧跟表格，保留结构）
            # 先创建空白段落
            blank_paras = [doc.add_paragraph("") for _ in range(self.blank_lines_after_table)]
            # 将空白段落移动到表格之后
            for idx, blank_para in enumerate(blank_paras):
                para_elem = blank_para._p
                para_elem.getparent().remove(para_elem)
                # 表格后第一个空白：索引=表格位置+1，第二个=表格位置+2
                insert_idx = doc._body._element.index(table_elem) + 1 + idx
                doc._body._element.insert(insert_idx, para_elem)
            
            self._log("  ✅ 表格后已添加2行空白内容（保留结构）")
            return True
        
        except Exception as e:
            self._log(f"  ❌ 表格创建失败：{str(e)}")
            import traceback
            self._log(f"  📝 详细错误：{traceback.format_exc()[:300]}")
            return False

    def _process_single_file(self, file_path):
        """处理单个docx文件（保留图片）"""
        try:
            file_name = os.path.basename(file_path)
            self._log(f"\n===== 处理文件：{file_name} =====")
            
            # 1. 备份原文件（防止数据丢失）
            backup_path = f"{file_path}.bak"
            shutil.copy2(file_path, backup_path)
            self._log(f"  📁 已备份原文件：{file_name}.bak")
            
            # 2. 检测文件名关键词
            keyword = self._check_filename_keyword(file_path)
            if not keyword:
                self._log(f"  ⚠️  文件名不含指定关键词，跳过处理")
                return "skip"
            
            self._log(f"  🔍 检测到关键词：{keyword}")
            
            # 3. 打开文档（使用原生方式，保留所有元素）
            doc = Document(file_path)
            
            # 4. 第二行插入表格（保留图片）
            create_success = self._insert_table_at_second_line(doc, keyword)
            
            # 5. 保存修改后的文档（安全保存，保留图片）
            doc.save(file_path)
            
            if create_success:
                self._log(f"  ✅ {file_name} 处理完成（图片已保留）")
                return "success"
            else:
                self._log(f"  ❌ {file_name} 处理失败（表格创建失败）")
                return "fail"
        
        except Exception as e:
            self._log(f"❌ 文件处理异常：{str(e)}")
            import traceback
            self._log(f"📝 详细错误：{traceback.format_exc()[:500]}")
            return "fail"

    def _batch_process(self):
        """批量处理文件夹下所有docx文件"""
        folder = self.folder_path.get()
        if not folder or not os.path.isdir(folder):
            messagebox.showerror("错误", "请选择有效的文件夹！")
            return
        
        # 清空日志
        self.log_text.delete(1.0, tk.END)
        self._log("🚀 开始批量处理docx文件（保留图片+第二行插入表格）...")
        self._log(f"📂 目标文件夹：{folder}")
        
        # 筛选所有docx文件
        docx_files = [
            os.path.join(folder, f) for f in os.listdir(folder)
            if f.lower().endswith(".docx") and os.path.isfile(os.path.join(folder, f))
        ]
        
        if not docx_files:
            self._log("⚠️  未找到任何.docx文件！")
            messagebox.showinfo("提示", "未找到任何.docx文件！")
            return
        
        self._log(f"📊 共找到 {len(docx_files)} 个docx文件")
        
        # 批量处理并统计结果
        success_count = 0
        fail_count = 0
        skip_count = 0
        for file_path in docx_files:
            result = self._process_single_file(file_path)
            if result == "success":
                success_count += 1
            elif result == "fail":
                fail_count += 1
            elif result == "skip":
                skip_count += 1
        
        # 处理完成统计提示
        result_msg = (
            f"\n✅ 批量处理完成！\n"
            f"✅ 成功添加表格（保留图片）：{success_count}个\n"
            f"❌ 处理失败：{fail_count}个\n"
            f"⚠️  无关键词跳过：{skip_count}个"
        )
        self._log(result_msg)
        messagebox.showinfo("处理完成", result_msg)

if __name__ == "__main__":
    # 适配tkinter中文显示
    root = tk.Tk()
    root.option_add("*Font", "SimHei 9")
    # 启动主程序
    app = DocxBatchTableTool(root)
    root.mainloop()
