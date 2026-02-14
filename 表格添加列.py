#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
用Python 3.8.7实现批量修改一个文件夹里面的.docx文件，并且生成一个GUI界面进行操作。
在.docx文件表格第三列右侧添加三列表格并且填充内容：
天线高度(cm)    天线极化    转台角度(deg)
200         H       ——
200         H       ——
200         H       ——
200         H       ——
200         H       ——
200         H       ——

原来的第四列表格和内容放在第7列



批量修改docx表格工具 - Python 3.8.7 + python-docx 0.8.11
功能：1. 第三列右侧加3列并填充指定内容 2. 原第四列移第七列 3. 显示完整表格边框
"""
import os
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn

class DocxTableModifier:
    def __init__(self, root):
        self.root = root
        self.root.title("Docx表格批量修改工具（显示完整边框）")
        self.root.geometry("700x550")
        
        # 新增3列的固定内容（表头+数据行）
        self.new_col_headers = ["天线高度(cm)", "天线极化", "转台角度(deg)"]
        self.new_col_data = ["200", "H", "——"]
        
        self.folder_path = tk.StringVar()
        self._create_gui()

    def _create_gui(self):
        """创建GUI界面"""
        # 1. 文件夹选择区域
        frame1 = tk.Frame(self.root, padx=10, pady=10)
        frame1.pack(fill=tk.X)
        
        tk.Label(frame1, text="目标文件夹：").pack(side=tk.LEFT)
        tk.Entry(frame1, textvariable=self.folder_path, width=50).pack(side=tk.LEFT, padx=5)
        tk.Button(frame1, text="选择文件夹", command=self._select_folder).pack(side=tk.LEFT)
        
        # 2. 执行按钮
        frame2 = tk.Frame(self.root, padx=10, pady=5)
        frame2.pack(fill=tk.X)
        
        tk.Button(
            frame2, text="开始批量处理", 
            command=self._batch_process,
            bg="#4CAF50", fg="white", font=("Arial", 10, "bold")
        ).pack(side=tk.LEFT)
        
        # 3. 日志显示区域
        frame3 = tk.Frame(self.root, padx=10, pady=10)
        frame3.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(frame3, text="处理日志：").pack(anchor=tk.W)
        self.log_text = scrolledtext.ScrolledText(frame3, height=20, font=("Consolas", 9))
        self.log_text.pack(fill=tk.BOTH, expand=True)

    def _select_folder(self):
        """选择目标文件夹"""
        folder = filedialog.askdirectory(title="选择包含docx文件的文件夹")
        if folder:
            self.folder_path.set(folder)
            self._log(f"已选择文件夹：{folder}")

    def _log(self, msg):
        """添加日志信息"""
        self.log_text.insert(tk.END, f"{msg}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def _set_cell_border(self, cell):
        """
        为单个单元格设置完整边框（低版本python-docx兼容方案）
        :param cell: 单元格对象
        """
        # 定义边框样式：黑色、0.5磅实线
        borders = ["top", "bottom", "left", "right"]
        for border_name in borders:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")       # 边框类型：实线
            border.set(qn("w:sz"), "4")             # 边框宽度：4=0.5磅（sz单位是1/8磅）
            border.set(qn("w:color"), "000000")     # 边框颜色：黑色
            border.set(qn("w:space"), "0")          # 边框间距：0
            cell._tc.get_or_add_tcPr().append(border)

    def _rebuild_table(self, table):
        """重建表格：读取原数据+构造新结构"""
        # 1. 读取原表格所有内容
        original_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            original_data.append(row_data)
        
        if not original_data:
            return None
        
        # 2. 构造新表格数据：原1-3列 + 新增3列 + 原4列（移第七列）
        new_table_data = []
        for idx, row in enumerate(original_data):
            # 补全原行数据（避免列数不足）
            row += [""] * (4 - len(row))
            
            # 表头行填新增列标题，数据行填固定内容
            new_cols = self.new_col_headers if idx == 0 else self.new_col_data
            
            # 新行结构：原1-3列 + 新增3列 + 原4列
            new_row = row[0:3] + new_cols + [row[3]]
            new_table_data.append(new_row)
        
        return new_table_data

    def _modify_docx_table(self, file_path):
        """修改单个docx文件的表格（含边框设置）"""
        # 1. 备份原文件
        backup_path = file_path + ".bak"
        shutil.copy2(file_path, backup_path)
        self._log(f"已备份原文件：{backup_path}")
        
        # 2. 打开文档并处理表格
        doc = Document(file_path)
        table_count = 0
        
        # 遍历所有表格，删除原表格并插入新表格
        tables_to_remove = []
        new_tables = []
        
        for table in doc.tables:
            table_count += 1
            self._log(f"  处理第{table_count}个表格（原行数：{len(table.rows)}）")
            
            # 跳过列数不足的表格
            if len(table.columns) < 4:
                self._log(f"  警告：第{table_count}个表格列数不足4列，跳过处理")
                continue
            
            # 3. 重建表格数据
            new_table_data = self._rebuild_table(table)
            if not new_table_data:
                self._log(f"  警告：第{table_count}个表格无数据，跳过处理")
                continue
            
            # 4. 记录原表格位置+删除原表格
            tables_to_remove.append(table)
            # 获取原表格的位置（在文档中的段落索引）
            table_paragraph = table._element.getparent()
            # 创建新表格
            new_table = doc.add_table(rows=len(new_table_data), cols=7)
            new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 调整表格列宽（可选，优化显示）
            for col in new_table.columns:
                col.width = Pt(60)  # 每列宽度60磅
            
            # 5. 填充新表格数据+设置边框
            for row_idx, row_data in enumerate(new_table_data):
                row_cells = new_table.rows[row_idx].cells
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(row_cells):
                        cell = row_cells[col_idx]
                        cell.text = cell_text
                        # 为每个单元格设置完整边框
                        self._set_cell_border(cell)
            
            new_tables.append((table_paragraph, new_table))
        
        # 6. 删除原表格+将新表格插入原位置
        for table in tables_to_remove:
            table._element.getparent().remove(table._element)
        for para, new_table in new_tables:
            para.addnext(new_table._element)
        
        # 7. 保存修改后的文档
        doc.save(file_path)
        self._log(f"已完成文件修改：{file_path}")
        return True

    def _batch_process(self):
        """批量处理文件夹下的docx文件"""
        folder = self.folder_path.get()
        if not folder or not os.path.exists(folder):
            messagebox.showerror("错误", "请选择有效的文件夹！")
            return
        
        # 清空日志
        self.log_text.delete(1.0, tk.END)
        self._log("开始批量处理...")
        
        # 遍历所有docx文件
        docx_files = [f for f in os.listdir(folder) if f.lower().endswith(".docx")]
        if not docx_files:
            self._log("未找到任何.docx文件！")
            messagebox.showinfo("提示", "未找到任何.docx文件！")
            return
        
        self._log(f"共找到{len(docx_files)}个docx文件，开始处理...")
        
        success_count = 0
        fail_count = 0
        
        for file_name in docx_files:
            file_path = os.path.join(folder, file_name)
            self._log(f"\n处理文件：{file_name}")
            
            try:
                self._modify_docx_table(file_path)
                success_count += 1
            except Exception as e:
                self._log(f"  处理失败：{str(e)}")
                fail_count += 1
        
        # 处理完成提示
        result_msg = f"处理完成！成功：{success_count}个，失败：{fail_count}个"
        self._log(f"\n{result_msg}")
        messagebox.showinfo("完成", result_msg)

if __name__ == "__main__":
    # 适配Python 3.8.7的tkinter中文显示
    root = tk.Tk()
    root.option_add("*Font", "SimHei 9")
    app = DocxTableModifier(root)
    root.mainloop()
