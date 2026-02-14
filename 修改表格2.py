'''
用Python 3.8.7实现批量修改一个文件夹里面的.docx文件，并且生成一个GUI界面进行操作。
删除“Test Report”；
删除每个文档中的第一个表格；
“Final_Result”替换为“试验结果图:”；
“Frequency”替换为“频率”；
“QuasiPeak”替换为“准峰值”；
“Margin”替换为“裕量”；
“Limit”替换为“限值”；
删除表格的第5列到第9列；
交换表格第3列和第4列的内容；

'''
import os
import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox
from docx import Document
import traceback

class DocxBatchProcessor:
    def __init__(self, root):
        self.root = root
        self.root.title("DOCX批量处理工具")
        self.root.geometry("800x600")
        
        # 选中的文件夹路径
        self.folder_path = tk.StringVar()
        
        # 创建GUI组件
        self._create_widgets()
        
    def _create_widgets(self):
        # 文件夹选择区域
        frame1 = tk.Frame(self.root, padx=10, pady=10)
        frame1.pack(fill=tk.X)
        
        tk.Label(frame1, text="目标文件夹:").pack(side=tk.LEFT)
        tk.Entry(frame1, textvariable=self.folder_path, width=60).pack(side=tk.LEFT, padx=5)
        tk.Button(frame1, text="选择文件夹", command=self.select_folder).pack(side=tk.LEFT)
        
        # 操作按钮区域
        frame2 = tk.Frame(self.root, padx=10, pady=5)
        frame2.pack(fill=tk.X)
        
        tk.Button(frame2, text="开始批量处理", command=self.process_documents, bg="#4CAF50", fg="white").pack(side=tk.LEFT, padx=5)
        tk.Button(frame2, text="清空日志", command=self.clear_log, bg="#f44336", fg="white").pack(side=tk.LEFT, padx=5)
        
        # 日志显示区域
        frame3 = tk.Frame(self.root, padx=10, pady=10)
        frame3.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(frame3, text="处理日志:").pack(anchor=tk.W)
        self.log_text = scrolledtext.ScrolledText(frame3, wrap=tk.WORD, height=20)
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
    def select_folder(self):
        """选择目标文件夹"""
        folder = filedialog.askdirectory(title="选择包含docx文件的文件夹")
        if folder:
            self.folder_path.set(folder)
            self.log(f"已选择文件夹: {folder}")
            
    def log(self, message):
        """添加日志信息"""
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
        
    def clear_log(self):
        """清空日志"""
        self.log_text.delete(1.0, tk.END)
        self.log("日志已清空")
        
    def process_single_document(self, file_path):
        """处理单个docx文件"""
        try:
            # 打开文档
            doc = Document(file_path)
            self.log(f"开始处理文件: {os.path.basename(file_path)}")
            
            # 1. 删除所有"Test Report"文本
            self.remove_text(doc, "Test Report")
            self.log("  - 已删除所有'Test Report'文本")
            
            # 2. 删除第一个表格
            if doc.tables:
                first_table = doc.tables[0]
                # 获取表格所在的段落并删除整个表格
                table_element = first_table._element
                table_element.getparent().remove(table_element)
                self.log("  - 已删除第一个表格")
            else:
                self.log("  - 文档中未找到表格，跳过删除第一个表格操作")
                
            # 3. 批量替换文本
            replace_pairs = {
                "Final_Result": "试验结果图:",
                "Frequency": "频率",
                "QuasiPeak": "准峰值",
                "Margin": "裕量",
                "Limit": "限值"
            }
            self.batch_replace_text(doc, replace_pairs)
            self.log("  - 已完成文本批量替换")
            
            # 4. 删除所有表格的第5列到第9列（索引从0开始，对应4-8）
            self.remove_table_columns(doc, start_col=4, end_col=8)
            self.log("  - 已删除所有表格的第5列到第9列")
            
            # 5. 新增功能：交换所有表格的第3列和第4列内容（索引2和3）
            self.swap_table_columns(doc, col1=2, col2=3)
            self.log("  - 已交换所有表格的第3列和第4列内容")
            
            # 保存修改后的文档（覆盖原文件）
            doc.save(file_path)
            self.log(f"  - 文件处理完成: {os.path.basename(file_path)}")
            return True
            
        except Exception as e:
            self.log(f"  - 处理文件出错: {str(e)}")
            self.log(f"  - 错误详情: {traceback.format_exc()}")
            return False
            
    def remove_text(self, doc, text_to_remove):
        """删除文档中指定文本"""
        # 遍历所有段落
        for para in doc.paragraphs:
            if text_to_remove in para.text:
                para.text = para.text.replace(text_to_remove, "")
        
        # 遍历所有表格中的单元格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if text_to_remove in cell.text:
                        cell.text = cell.text.replace(text_to_remove, "")
                        
    def batch_replace_text(self, doc, replace_pairs):
        """批量替换文本"""
        # 替换段落中的文本
        for para in doc.paragraphs:
            for old_text, new_text in replace_pairs.items():
                if old_text in para.text:
                    para.text = para.text.replace(old_text, new_text)
        
        # 替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old_text, new_text in replace_pairs.items():
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)
                            
    def remove_table_columns(self, doc, start_col, end_col):
        """删除表格中指定范围的列（索引从0开始）"""
        for table in doc.tables:
            # 获取表格的最大列数
            max_cols = max(len(row.cells) for row in table.rows)
            if start_col >= max_cols:
                self.log(f"  - 表格列数不足，跳过列删除操作（当前最大列数: {max_cols}）")
                continue
                
            # 调整结束列索引，避免越界
            actual_end_col = min(end_col, max_cols - 1)
            
            # 从后往前删除列（避免索引错乱）
            for col_idx in range(actual_end_col, start_col - 1, -1):
                for row in table.rows:
                    if len(row.cells) > col_idx:
                        cell = row.cells[col_idx]
                        cell._element.getparent().remove(cell._element)
    
    def swap_table_columns(self, doc, col1, col2):
        """
        交换表格中指定两列的内容
        :param doc: Document对象
        :param col1: 第一列索引（从0开始）
        :param col2: 第二列索引（从0开始）
        """
        for table in doc.tables:
            # 获取表格的最大列数
            max_cols = max(len(row.cells) for row in table.rows)
            
            # 检查列索引是否有效
            if col1 >= max_cols or col2 >= max_cols:
                self.log(f"  - 表格列数不足（当前最大列数: {max_cols}），跳过列交换操作")
                continue
            
            # 遍历每一行，交换指定列的内容
            for row in table.rows:
                # 确保当前行有足够的列
                if len(row.cells) > max(col1, col2):
                    # 暂存第一列内容，避免覆盖
                    temp_text = row.cells[col1].text
                    # 交换内容
                    row.cells[col1].text = row.cells[col2].text
                    row.cells[col2].text = temp_text

    def process_documents(self):
        """批量处理文件夹中的所有docx文件"""
        folder = self.folder_path.get()
        if not folder or not os.path.exists(folder):
            messagebox.showerror("错误", "请选择有效的文件夹！")
            return
            
        # 获取所有docx文件
        docx_files = [f for f in os.listdir(folder) if f.lower().endswith(".docx")]
        if not docx_files:
            messagebox.showinfo("提示", "文件夹中未找到docx文件！")
            return
            
        self.log(f"找到 {len(docx_files)} 个docx文件，开始批量处理...")
        
        success_count = 0
        fail_count = 0
        
        for filename in docx_files:
            file_path = os.path.join(folder, filename)
            if self.process_single_document(file_path):
                success_count += 1
            else:
                fail_count += 1
        
        # 处理完成统计
        self.log("="*50)
        self.log(f"批量处理完成！成功: {success_count} 个，失败: {fail_count} 个")
        messagebox.showinfo("完成", f"批量处理完成！\n成功: {success_count} 个\n失败: {fail_count} 个")

if __name__ == "__main__":
    # 安装依赖提示（首次运行前需要执行）
    print("提示：首次运行请先安装依赖：pip install python-docx")
    
    # 创建并运行GUI
    root = tk.Tk()
    app = DocxBatchProcessor(root)
    root.mainloop()
